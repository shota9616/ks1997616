#!/usr/bin/env python3
"""
省力化補助金（一般型）申請書類生成スクリプト v10.5 完全版
採択レベルの事業計画書を自動生成

【対応機能】
1. ヒアリングシート（10シート+財務情報）からの完全なデータ読み込み
2. 事業者概要ネストテーブルの完全入力（経営理念、経営戦略、直近実績など）
3. PREP法による散文形式の文章生成（各セクション600字以上）
4. SWOT分析の自動生成
5. 直近3年実績の表形式出力
6. nano-banana-pro-preview による高品質図解生成（11種類）
7. 全11種類の書類テンプレート対応

【使用方法】
python scripts/main.py --hearing ヒアリングシート.xlsx --output ./output --template-dir ./templates
"""

import os
import sys
import shutil
import base64
import time
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple

import openpyxl
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Gemini API
try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False


# =============================================================================
# 設定値（Phase 2: 外部化）
# =============================================================================

class Config:
    """ハードコード値を集約した設定クラス"""
    # Gemini API
    GEMINI_MODEL = "nano-banana-pro-preview"
    GEMINI_RETRY_MAX = 3
    GEMINI_RETRY_BASE_DELAY = 2  # seconds
    GEMINI_INTER_REQUEST_DELAY = 2  # seconds

    # 人件費・稼働
    HOURLY_WAGE = 2500  # 円
    WORKING_DAYS_PER_YEAR = 250
    WORKING_DAYS_PER_MONTH = 22

    # 財務推計（デフォルト値）
    DEPRECIATION_YEARS = 5
    GROWTH_RATE = 1.05  # 付加価値額の年間成長率（公募要領: 年率4%以上、余裕持ち5%）
    SALARY_GROWTH_RATE = 1.025  # 給与支給総額の年間成長率（公募要領: 年率2%以上、余裕持ち2.5%）
    PROFIT_GROWTH_RATE = 1.05  # 営業利益の年間成長率

    # 自動修正ループ用: デフォルト値（リセット用の不変定数）
    _GROWTH_RATE_DEFAULT = 1.05
    _SALARY_GROWTH_RATE_DEFAULT = 1.025
    _PROFIT_GROWTH_RATE_DEFAULT = 1.05

    @classmethod
    def reset_rates(cls):
        """自動修正ループ開始前に成長率をデフォルトにリセットする"""
        cls.GROWTH_RATE = cls._GROWTH_RATE_DEFAULT
        cls.SALARY_GROWTH_RATE = cls._SALARY_GROWTH_RATE_DEFAULT
        cls.PROFIT_GROWTH_RATE = cls._PROFIT_GROWTH_RATE_DEFAULT
    LABOR_COST_RATIO = 0.35  # 売上高に対する人件費比率
    SALARY_RATIO = 0.3  # 売上高に対する給与比率

    # デフォルト財務値
    DEFAULT_REVENUE = 50000000
    DEFAULT_PROFIT = 7000000

    # 業種別デフォルト有効求人倍率
    INDUSTRY_JOB_RATIOS = {
        "建設": 5.8,
        "建築": 5.8,
        "製造": 2.1,
        "IT": 3.9,
        "情報": 3.9,
        "飲食": 3.2,
        "サービス": 3.0,
        "小売": 2.4,
        "介護": 3.8,
        "運輸": 2.7,
    }
    DEFAULT_JOB_RATIO = 5.8

    # 業種別経営理念テンプレート（Phase 3）
    INDUSTRY_PHILOSOPHY_TEMPLATES = {
        "建設": "お客様の理想の住まいを実現し、地域に根ざした建築サービスを通じて社会に貢献する。安全で高品質な施工により、地域の発展とお客様の豊かな暮らしの実現に寄与することを使命とする。",
        "製造": "ものづくりの技術と品質にこだわり、顧客に信頼される製品を提供し続ける。生産性の向上と技術革新を通じて、日本のものづくり産業の発展に貢献することを使命とする。",
        "IT": "テクノロジーの力で社会課題を解決し、顧客のデジタル変革を支援する。最先端技術の活用と高い専門性により、持続可能な社会の実現に貢献することを使命とする。",
        "飲食": "安全で美味しい食を提供し、お客様の笑顔と健康に貢献する。地域の食文化を大切にしながら、従業員が誇りを持って働ける職場づくりを使命とする。",
        "サービス": "お客様一人ひとりに寄り添い、期待を超えるサービスを提供する。人と人とのつながりを大切に、地域社会に貢献することを使命とする。",
        "小売": "お客様に必要な商品を適正な価格で提供し、地域の暮らしを支える。品揃えと接客の質にこだわり、地域になくてはならない存在を目指すことを使命とする。",
    }
    DEFAULT_PHILOSOPHY_TEMPLATE = "お客様の理想の住まいを実現し、地域に根ざしたサービスを通じて社会に貢献する。{industry}における専門性を活かし、高品質なサービスで顧客満足と地域発展に寄与することを使命とする。"


# =============================================================================
# データクラス定義
# =============================================================================

@dataclass
class CompanyInfo:
    """企業基本情報"""
    name: str = ""
    representative: str = ""
    address: str = ""
    prefecture: str = ""
    postal_code: str = ""
    phone: str = ""
    established_date: str = ""
    capital: int = 0
    industry: str = ""
    business_description: str = ""
    employee_count: int = 0
    officer_count: int = 1
    url: str = ""
    # 財務情報
    revenue_2022: int = 0
    revenue_2023: int = 0
    revenue_2024: int = 0
    gross_profit_2022: int = 0
    gross_profit_2023: int = 0
    gross_profit_2024: int = 0
    operating_profit_2022: int = 0
    operating_profit_2023: int = 0
    operating_profit_2024: int = 0
    # 付加価値額算出用（決算書PDFから取得）
    labor_cost: int = 0           # 人件費合計
    depreciation: int = 0         # 減価償却費
    total_salary: int = 0         # 給与支給総額（役員報酬除く）


@dataclass
class LaborShortageInfo:
    """人手不足情報"""
    shortage_tasks: str = ""
    recruitment_period: str = ""
    applications: int = 0
    hired: int = 0
    overtime_hours: float = 0
    current_workers: int = 0
    desired_workers: int = 0
    job_openings_ratio: float = 0


@dataclass
class LaborSavingInfo:
    """省力化効果情報"""
    target_tasks: str = ""
    current_hours: float = 0
    target_hours: float = 0
    reduction_hours: float = 0
    reduction_rate: float = 0


@dataclass
class EquipmentInfo:
    """導入設備情報"""
    name: str = ""
    category: str = ""
    manufacturer: str = ""
    model: str = ""
    quantity: int = 1
    total_price: int = 0
    vendor: str = ""
    features: str = ""
    catalog_number: str = ""


@dataclass
class FundingInfo:
    """資金調達情報"""
    subsidy_amount: int = 0
    self_funding: int = 0
    total_investment: int = 0
    implementation_manager: str = ""
    implementation_period: str = ""
    bank_name: str = ""


@dataclass
class WorkProcess:
    """作業工程"""
    name: str = ""
    time_minutes: int = 0
    description: str = ""


@dataclass
class OfficerInfo:
    """役員情報"""
    name: str = ""
    position: str = ""
    birth_date: str = ""


@dataclass
class EmployeeInfo:
    """従業員情報"""
    name: str = ""
    birth_date: str = ""
    hire_date: str = ""


@dataclass
class ShareholderInfo:
    """株主情報"""
    name: str = ""
    shares: int = 0


@dataclass
class HearingData:
    """ヒアリングデータ全体"""
    company: CompanyInfo = field(default_factory=CompanyInfo)
    labor_shortage: LaborShortageInfo = field(default_factory=LaborShortageInfo)
    labor_saving: LaborSavingInfo = field(default_factory=LaborSavingInfo)
    equipment: EquipmentInfo = field(default_factory=EquipmentInfo)
    funding: FundingInfo = field(default_factory=FundingInfo)
    officers: List[OfficerInfo] = field(default_factory=list)
    employees: List[EmployeeInfo] = field(default_factory=list)
    shareholders: List[ShareholderInfo] = field(default_factory=list)
    before_processes: List[WorkProcess] = field(default_factory=list)
    after_processes: List[WorkProcess] = field(default_factory=list)
    # Phase 4: 追加フィールド
    motivation_background: str = ""  # なぜ今必要か（シート3）
    time_utilization_plan: str = ""  # 効果の活用計画（シート6）
    wage_increase_rate: float = 0.0  # 賃上げ率（シート7）
    wage_increase_target: str = ""  # 賃上げ対象者（シート7）
    wage_increase_timing: str = ""  # 賃上げ実施時期（シート7）


# =============================================================================
# ヒアリングシート読み込み
# =============================================================================

def _split_name(name: str) -> Tuple[str, str]:
    """姓名を分割する（全角スペース対応）（Phase 6）"""
    name = str(name).strip()
    # 全角スペース → 半角スペースに正規化してから分割
    normalized = name.replace('\u3000', ' ')
    parts = normalized.split()
    if len(parts) >= 2:
        return parts[0], ' '.join(parts[1:])
    return name, ''


def _find_sheet_in_workbook(wb, patterns: list):
    """ファジーマッチングでワークブック内のシートを検索する（Phase 6）"""
    for name in wb.sheetnames:
        for p in patterns:
            if p in name:
                return wb[name]
    return None


def read_hearing_sheet(file_path: str) -> HearingData:
    """ヒアリングシートから全データを読み込む"""
    print(f"📖 ヒアリングシート読み込み中: {file_path}")

    # Phase 1: リソースリーク防止 — try-except-finally で wb.close() を保証
    wb = None
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
    except FileNotFoundError:
        print(f"  ❌ ファイルが見つかりません: {file_path}")
        raise
    except Exception as ex:
        print(f"  ❌ ファイルオープンエラー: {ex}")
        raise

    data = HearingData()

    try:
        # ヘルパー関数
        def find_value(ws, labels, offset=1):
            """ラベルに対応する値を検索"""
            if isinstance(labels, str):
                labels = [labels]
            for row in range(1, 50):
                for col in range(1, 10):
                    val = ws.cell(row=row, column=col).value
                    if val:
                        for label in labels:
                            if label in str(val):
                                result = ws.cell(row=row, column=col + offset).value
                                # Phase 1: None ガード
                                return result if result is not None else ""
            return ""

        def find_int(ws, labels, offset=1, default=0):
            """整数値を検索"""
            val = find_value(ws, labels, offset)
            if not val:
                return default
            try:
                return int(float(str(val).replace(",", "").replace("円", "").replace("人", "").replace("名", "")))
            except (ValueError, TypeError):
                return default

        def find_float(ws, labels, offset=1, default=0.0):
            """浮動小数点値を検索"""
            val = find_value(ws, labels, offset)
            if not val:
                return default
            try:
                return float(str(val).replace(",", "").replace("%", "").replace("時間", ""))
            except (ValueError, TypeError):
                return default

        def find_sheet(patterns):
            """パターンに一致するシートを検索（Phase 6: ファジーマッチング関数を使用）"""
            return _find_sheet_in_workbook(wb, patterns)

        # ----- 1. 企業基本情報 -----
        ws = find_sheet(["企業基本情報", "1_"])
        if ws:
            data.company.name = str(find_value(ws, ["会社名", "法人名"]))
            data.company.representative = str(find_value(ws, ["代表者名", "代表取締役"]))
            data.company.prefecture = str(find_value(ws, ["都道府県"]))
            data.company.address = str(find_value(ws, ["市区町村"]))
            data.company.phone = str(find_value(ws, ["電話番号"]))
            data.company.industry = str(find_value(ws, ["業種"]))
            data.company.business_description = str(find_value(ws, ["事業内容", "主な事業内容"]))
            data.company.employee_count = find_int(ws, ["従業員数", "常勤"])
            data.company.established_date = str(find_value(ws, ["設立"]))
            data.company.capital = find_int(ws, ["資本金"])
            data.company.url = str(find_value(ws, ["URL", "ホームページ"]))

        # ----- 2. 人手不足の実態 -----
        ws = find_sheet(["人手不足", "2_"])
        if ws:
            data.labor_shortage.shortage_tasks = str(find_value(ws, ["困っている作業", "人が足りなくて"]))
            data.labor_shortage.recruitment_period = str(find_value(ws, ["求人を出している期間"]))
            data.labor_shortage.applications = find_int(ws, ["応募"])
            data.labor_shortage.hired = find_int(ws, ["採用できた"])
            data.labor_shortage.overtime_hours = find_float(ws, ["残業"])
            data.labor_shortage.current_workers = find_int(ws, ["何人でやっています"])
            data.labor_shortage.desired_workers = find_int(ws, ["何人いれば"])
            data.labor_shortage.job_openings_ratio = find_float(ws, ["有効求人倍率"])

        # ----- 3. 省力化効果 -----
        ws = find_sheet(["省力化効果", "4_"])
        if ws:
            data.labor_saving.target_tasks = str(find_value(ws, ["対象となる作業"]))
            data.labor_saving.current_hours = find_float(ws, ["導入前", "1日に何時間", "何時間その作業", "現在の作業時間", "今かかっている"])
            data.labor_saving.target_hours = find_float(ws, ["導入後", "導入したら何時間", "目標時間", "何時間になりそう", "短縮後"])
            if data.labor_saving.current_hours > 0:
                # Phase 1: マイナス削減率防止
                data.labor_saving.reduction_hours = max(0, data.labor_saving.current_hours - data.labor_saving.target_hours)
                data.labor_saving.reduction_rate = max(0, (data.labor_saving.reduction_hours / data.labor_saving.current_hours) * 100)
                if data.labor_saving.target_hours > data.labor_saving.current_hours:
                    print(f"  ⚠️ 警告: 導入後の時間({data.labor_saving.target_hours}h)が導入前({data.labor_saving.current_hours}h)より大きいです。削減率を0%にしました。")

        # ----- 4. 導入設備 -----
        ws = find_sheet(["5_導入", "導入設備"])
        if ws:
            data.equipment.name = str(find_value(ws, ["設備の名前", "設備名"]))
            data.equipment.category = str(find_value(ws, ["設備カテゴリ", "何をするもの"]))
            data.equipment.manufacturer = str(find_value(ws, ["メーカー"]))
            data.equipment.model = str(find_value(ws, ["型番"]))
            data.equipment.quantity = find_int(ws, ["数量"], default=1)
            data.equipment.vendor = str(find_value(ws, ["購入先", "ベンダー", "どこから買"]))
            data.equipment.total_price = find_int(ws, ["金額", "税抜", "いくら"])
            data.equipment.features = str(find_value(ws, ["カスタマイズ", "特徴"]))
            data.equipment.catalog_number = str(find_value(ws, ["カタログ", "登録番号"]))

        # ----- 5. 資金調達・体制 -----
        ws = find_sheet(["資金調達", "体制", "8_"])
        if ws:
            data.funding.total_investment = find_int(ws, ["投資総額", "設備の金額"])
            data.funding.subsidy_amount = find_int(ws, ["補助金申請額", "補助金"])
            data.funding.self_funding = find_int(ws, ["自己資金"])
            data.funding.implementation_manager = str(find_value(ws, ["責任者"]))
            data.funding.implementation_period = str(find_value(ws, ["実施期間", "導入時期"]))
            data.funding.bank_name = str(find_value(ws, ["銀行", "借入先", "取引銀行"]))

        if data.funding.total_investment == 0:
            data.funding.total_investment = data.equipment.total_price

        # ----- 6. 財務情報 -----
        ws = find_sheet(["11_財務", "財務情報"])
        if ws:
            base_revenue = find_int(ws, ["売上高"], default=Config.DEFAULT_REVENUE)
            base_gross_profit = find_int(ws, ["売上総利益", "粗利"], default=int(base_revenue * 0.7))
            base_profit = find_int(ws, ["営業利益"], default=Config.DEFAULT_PROFIT)

            # 基準年度から3年分を推計（Phase 2: Config参照）
            data.company.revenue_2024 = base_revenue
            data.company.revenue_2023 = int(base_revenue / Config.GROWTH_RATE)
            data.company.revenue_2022 = int(base_revenue / Config.GROWTH_RATE / Config.GROWTH_RATE)

            data.company.gross_profit_2024 = base_gross_profit
            data.company.gross_profit_2023 = int(base_gross_profit / Config.GROWTH_RATE)
            data.company.gross_profit_2022 = int(base_gross_profit / Config.GROWTH_RATE / Config.GROWTH_RATE)

            data.company.operating_profit_2024 = base_profit
            data.company.operating_profit_2023 = int(base_profit / Config.PROFIT_GROWTH_RATE)
            data.company.operating_profit_2022 = int(base_profit / Config.PROFIT_GROWTH_RATE / Config.PROFIT_GROWTH_RATE)
        else:
            # デフォルト値
            data.company.revenue_2022, data.company.revenue_2023, data.company.revenue_2024 = 47000000, 49000000, 50000000
            data.company.gross_profit_2022, data.company.gross_profit_2023, data.company.gross_profit_2024 = 33000000, 34000000, 35000000
            data.company.operating_profit_2022, data.company.operating_profit_2023, data.company.operating_profit_2024 = 6000000, 6500000, 7000000

        # ----- 7. 役員情報（株主セクションの手前まで）-----
        ws = find_sheet(["10_役員", "役員_株主"])
        if ws:
            in_officer_section = False
            for row in range(1, ws.max_row + 1):
                col1 = ws.cell(row=row, column=1).value
                name = ws.cell(row=row, column=2).value

                if col1 and "役員情報" in str(col1):
                    in_officer_section = True
                    continue
                if col1 and "株主情報" in str(col1):
                    break

                if in_officer_section and name and str(name).strip() and "氏名" not in str(name):
                    data.officers.append(OfficerInfo(
                        name=str(name),
                        position=str(ws.cell(row=row, column=3).value or "役員"),
                        birth_date=str(ws.cell(row=row, column=4).value or "")
                    ))
            data.company.officer_count = max(len(data.officers), 1)

        # ----- 8. 従業員情報 -----
        ws = find_sheet(["9_従業員", "従業員情報"])
        if ws:
            for row in range(2, ws.max_row + 1):
                name = ws.cell(row=row, column=2).value
                if name and str(name).strip() and "氏名" not in str(name):
                    data.employees.append(EmployeeInfo(
                        name=str(name),
                        birth_date=str(ws.cell(row=row, column=3).value or ""),
                        hire_date=str(ws.cell(row=row, column=4).value or "")
                    ))

        # ----- 9. 株主情報 -----
        ws = find_sheet(["10_役員", "役員_株主"])
        if ws:
            in_shareholder_section = False
            for row in range(1, ws.max_row + 1):
                col1 = ws.cell(row=row, column=1).value
                name = ws.cell(row=row, column=2).value

                if col1 and "株主情報" in str(col1):
                    in_shareholder_section = True
                    continue

                if in_shareholder_section and name and str(name).strip() and "株主名" not in str(name):
                    try:
                        shares = int(ws.cell(row=row, column=3).value or 0)
                    except (ValueError, TypeError):
                        shares = 0
                    data.shareholders.append(ShareholderInfo(name=str(name), shares=shares))

        # ----- Phase 4: シート3（なぜ今必要か）-----
        ws = find_sheet(["なぜ今", "3_"])
        if ws:
            data.motivation_background = str(find_value(ws, ["背景", "なぜ今", "理由", "きっかけ"]))

        # ----- Phase 4: シート6（効果の活用計画）-----
        ws = find_sheet(["効果の活用", "6_"])
        if ws:
            data.time_utilization_plan = str(find_value(ws, ["活用", "浮いた時間", "時間の使い方"]))

        # ----- Phase 4: シート7（賃上げ計画）-----
        ws = find_sheet(["賃上げ", "7_"])
        if ws:
            data.wage_increase_rate = find_float(ws, ["賃上げ率", "引上げ率"])
            data.wage_increase_target = str(find_value(ws, ["対象者", "対象"]))
            data.wage_increase_timing = str(find_value(ws, ["実施時期", "いつから"]))

        # ----- 10. 工程データ生成 -----
        data.before_processes, data.after_processes = generate_processes(data)

    finally:
        # Phase 1: リソースリーク防止 — 必ず wb.close() を実行
        if wb is not None:
            wb.close()

    # フォールバック計算（0値の補完）
    if data.labor_saving.current_hours == 0 and data.before_processes:
        data.labor_saving.current_hours = sum(p.time_minutes for p in data.before_processes) / 60
        print(f"  ⚠️ current_hours=0 → before_processesから推計: {data.labor_saving.current_hours:.1f}h")
    if data.labor_saving.target_hours == 0 and data.after_processes:
        data.labor_saving.target_hours = sum(p.time_minutes for p in data.after_processes) / 60
        print(f"  ⚠️ target_hours=0 → after_processesから推計: {data.labor_saving.target_hours:.1f}h")
    if data.labor_saving.current_hours > 0 and data.labor_saving.target_hours > 0:
        data.labor_saving.reduction_hours = max(0, data.labor_saving.current_hours - data.labor_saving.target_hours)
        data.labor_saving.reduction_rate = max(0, (data.labor_saving.reduction_hours / data.labor_saving.current_hours) * 100)
    if data.funding.subsidy_amount == 0 and data.funding.total_investment > 0:
        data.funding.subsidy_amount = int(data.funding.total_investment * 0.5)
        print(f"  ⚠️ subsidy_amount=0 → total_investment×0.5で推計: {data.funding.subsidy_amount:,}円")
    if data.funding.self_funding == 0 and data.funding.total_investment > 0:
        data.funding.self_funding = data.funding.total_investment - data.funding.subsidy_amount
        print(f"  ⚠️ self_funding=0 → 差額で推計: {data.funding.self_funding:,}円")

    # 読み込み結果表示
    print(f"  ✅ 企業名: {data.company.name}")
    print(f"  ✅ 業種: {data.company.industry}")
    print(f"  ✅ 役員: {data.company.officer_count}名 / 従業員: {data.company.employee_count}名")
    print(f"  ✅ 設備: {data.equipment.name}")
    print(f"  ✅ 投資額: {data.equipment.total_price:,}円")
    print(f"  ✅ 補助金: {data.funding.subsidy_amount:,}円")
    print(f"  ✅ 削減率: {data.labor_saving.reduction_rate:.1f}%")
    print(f"  ✅ 売上高(2024): {data.company.revenue_2024:,}円")

    return data


def validate_hearing_data(data: HearingData) -> List[str]:
    """ヒアリングデータの必須フィールドを検証し、問題リストを返す"""
    issues = []
    if data.labor_saving.current_hours <= 0:
        issues.append("導入前の作業時間(current_hours)が0です")
    if data.labor_saving.target_hours <= 0:
        issues.append("導入後の作業時間(target_hours)が0です")
    if data.funding.subsidy_amount <= 0:
        issues.append("補助金申請額(subsidy_amount)が0です")
    if data.funding.total_investment <= 0:
        issues.append("投資総額(total_investment)が0です")
    if data.equipment.total_price <= 0:
        issues.append("設備価格(total_price)が0です")
    if not data.company.name or data.company.name.strip() == "":
        issues.append("企業名が空です")
    if data.company.employee_count <= 0:
        issues.append("従業員数(employee_count)が0です")
    return issues


def generate_processes(data: HearingData) -> Tuple[List[WorkProcess], List[WorkProcess]]:
    """業種に応じた工程データを生成（Phase 3: 6業種対応）"""
    industry = data.company.industry

    if "建設" in industry or "建築" in industry:
        before = [
            WorkProcess("顧客打合せ", 60, "要件ヒアリング"),
            WorkProcess("図面作成", 120, "CAD設計"),
            WorkProcess("数量拾い出し", 90, "手作業計算"),
            WorkProcess("単価確認", 120, "見積依頼"),
            WorkProcess("見積書作成", 60, "書類作成"),
            WorkProcess("顧客説明", 30, "提案"),
        ]
        after = [
            WorkProcess("顧客打合せ", 60, "要件ヒアリング"),
            WorkProcess("図面作成", 120, "CAD設計"),
            WorkProcess("数量拾い出し", 10, "AI自動計算"),
            WorkProcess("単価確認", 15, "AIマッチング"),
            WorkProcess("見積書作成", 10, "自動生成"),
            WorkProcess("顧客説明", 30, "提案"),
        ]
    elif "製造" in industry:
        before = [
            WorkProcess("受注処理", 30, "注文確認・伝票起票"),
            WorkProcess("生産計画", 45, "手動スケジューリング"),
            WorkProcess("部材手配", 40, "在庫確認・発注"),
            WorkProcess("加工", 60, "手動作業"),
            WorkProcess("検品", 45, "目視確認"),
            WorkProcess("出荷準備", 30, "梱包・伝票作成"),
        ]
        after = [
            WorkProcess("受注処理", 10, "自動取り込み"),
            WorkProcess("生産計画", 10, "AI最適化"),
            WorkProcess("部材手配", 10, "自動発注"),
            WorkProcess("加工", 30, "自動化"),
            WorkProcess("検品", 15, "AI検査"),
            WorkProcess("出荷準備", 15, "自動梱包"),
        ]
    elif "IT" in industry or "情報" in industry:
        before = [
            WorkProcess("要件定義", 60, "顧客ヒアリング"),
            WorkProcess("設計", 90, "手動設計書作成"),
            WorkProcess("コーディング", 120, "手動開発"),
            WorkProcess("テスト", 60, "手動テスト"),
            WorkProcess("ドキュメント作成", 45, "手動作成"),
            WorkProcess("デプロイ", 30, "手動デプロイ"),
        ]
        after = [
            WorkProcess("要件定義", 60, "顧客ヒアリング"),
            WorkProcess("設計", 30, "AI支援設計"),
            WorkProcess("コーディング", 40, "AI支援開発"),
            WorkProcess("テスト", 15, "自動テスト"),
            WorkProcess("ドキュメント作成", 10, "自動生成"),
            WorkProcess("デプロイ", 10, "自動デプロイ"),
        ]
    elif "飲食" in industry:
        before = [
            WorkProcess("食材発注", 30, "在庫確認・手動発注"),
            WorkProcess("仕込み", 60, "手作業調理"),
            WorkProcess("注文受付", 20, "口頭・手書き"),
            WorkProcess("調理", 45, "手作業調理"),
            WorkProcess("会計", 15, "手動レジ"),
            WorkProcess("在庫管理", 30, "手動棚卸し"),
        ]
        after = [
            WorkProcess("食材発注", 5, "AI自動発注"),
            WorkProcess("仕込み", 40, "一部自動化"),
            WorkProcess("注文受付", 5, "タブレット注文"),
            WorkProcess("調理", 30, "調理支援機器"),
            WorkProcess("会計", 5, "自動精算"),
            WorkProcess("在庫管理", 5, "自動管理"),
        ]
    elif "サービス" in industry or "介護" in industry:
        before = [
            WorkProcess("予約管理", 30, "手動台帳管理"),
            WorkProcess("顧客対応", 45, "電話・来客対応"),
            WorkProcess("書類作成", 40, "手動作成"),
            WorkProcess("実作業", 60, "手作業"),
            WorkProcess("報告書作成", 30, "手書き"),
            WorkProcess("請求処理", 25, "手動計算"),
        ]
        after = [
            WorkProcess("予約管理", 5, "オンライン自動管理"),
            WorkProcess("顧客対応", 20, "AI自動応答併用"),
            WorkProcess("書類作成", 10, "自動生成"),
            WorkProcess("実作業", 40, "機器支援"),
            WorkProcess("報告書作成", 5, "自動生成"),
            WorkProcess("請求処理", 5, "自動計算"),
        ]
    elif "小売" in industry:
        before = [
            WorkProcess("発注業務", 30, "手動発注・在庫確認"),
            WorkProcess("検品", 25, "目視確認"),
            WorkProcess("陳列", 30, "手作業"),
            WorkProcess("接客", 40, "対面対応"),
            WorkProcess("会計", 20, "手動レジ"),
            WorkProcess("棚卸し", 45, "手動カウント"),
        ]
        after = [
            WorkProcess("発注業務", 5, "AI自動発注"),
            WorkProcess("検品", 10, "バーコード自動検品"),
            WorkProcess("陳列", 20, "最適配置提案"),
            WorkProcess("接客", 30, "セルフ+有人併用"),
            WorkProcess("会計", 5, "セルフレジ"),
            WorkProcess("棚卸し", 10, "自動在庫管理"),
        ]
    else:
        # デフォルト（汎用）
        before = [
            WorkProcess("検査", 30, "品質確認"),
            WorkProcess("準備", 20, "セットアップ"),
            WorkProcess("加工", 60, "手動作業"),
            WorkProcess("検品", 45, "目視確認"),
            WorkProcess("仕上げ", 30, "調整"),
            WorkProcess("梱包", 25, "出荷準備"),
        ]
        after = [
            WorkProcess("検査", 10, "自動検査"),
            WorkProcess("準備", 15, "自動セット"),
            WorkProcess("加工", 30, "自動化"),
            WorkProcess("検品", 15, "AI検査"),
            WorkProcess("仕上げ", 20, "効率化"),
            WorkProcess("梱包", 20, "効率化"),
        ]
    return before, after


# =============================================================================
# PREP法による文章生成
# =============================================================================

class ContentGenerator:
    """採択レベルの文章を生成するクラス"""

    def __init__(self, data: HearingData):
        self.data = data
        self.c = data.company
        self.s = data.labor_shortage
        self.l = data.labor_saving
        self.e = data.equipment
        self.f = data.funding
        # Phase 2: Config参照 + _get_default_job_ratio メソッド
        self.job_ratio = self.s.job_openings_ratio if self.s.job_openings_ratio > 0 else self._get_default_job_ratio()
        self.manufacturer = self.e.manufacturer if self.e.manufacturer else "オーダーメイド開発"
        self.model = self.e.model if self.e.model else "カスタム仕様"

    def _get_default_job_ratio(self) -> float:
        """業種別デフォルト有効求人倍率を取得（Phase 2）"""
        for keyword, ratio in Config.INDUSTRY_JOB_RATIOS.items():
            if keyword in self.c.industry:
                return ratio
        return Config.DEFAULT_JOB_RATIO

    def _get_industry_philosophy(self) -> str:
        """業種別経営理念テンプレートを取得（Phase 3）"""
        for keyword, template in Config.INDUSTRY_PHILOSOPHY_TEMPLATES.items():
            if keyword in self.c.industry:
                return template
        return Config.DEFAULT_PHILOSOPHY_TEMPLATE.format(industry=self.c.industry)

    def generate_business_overview_table_data(self) -> dict:
        """事業者概要テーブル用のデータを生成"""
        return {
            "事業者名": self.c.name,
            # Phase 3: 業種別経営理念テンプレート
            "経営理念": self._get_industry_philosophy(),
            "経営戦略": f"{self.c.industry}として、{self.c.business_description}を専門に、高品質なサービスで顧客満足を追求。デジタル化・AI活用による業務効率向上で競争力を強化し、限られた人員で最大の成果を創出する戦略を推進。",
            "事業コンセプト": f"対象エリア：{self.c.prefecture}を中心とした地域。ターゲット：{self.c.industry}サービスを必要とする個人・法人顧客。提供サービス：{self.c.business_description}。強み：専門技術と豊富な経験に基づく高品質サービス。",
            "事業内容": f"①{self.c.business_description}の提供\n②差別化ポイント：専門資格者による高品質サービス、地域特性への深い理解\n③顧客価値：専門性の高いサービス提供、迅速な対応、長期的な信頼関係構築",
            "長期的なビジョン": f"5年後：{self.e.name}の活用による業務効率化を完了し、受注能力を1.5倍に拡大。従業員の働き方改革を実現。10年後：{self.c.prefecture}地域でトップクラスの{self.c.industry}事業者を目指し、後継者育成と事業承継の基盤を確立する。",
            "直近実績": {
                "売上金額": [self.c.revenue_2022, self.c.revenue_2023, self.c.revenue_2024],
                "売上総利益": [self.c.gross_profit_2022, self.c.gross_profit_2023, self.c.gross_profit_2024],
                "営業利益": [self.c.operating_profit_2022, self.c.operating_profit_2023, self.c.operating_profit_2024],
                "従業員数": [self.c.employee_count, self.c.employee_count, self.c.employee_count],
            }
        }

    def generate_section_1_1(self) -> str:
        """1-1 現状分析（PREP法、600字以上）"""
        added_value_2024 = self.c.operating_profit_2024 + int(self.c.revenue_2024 * Config.LABOR_COST_RATIO) + self.c.depreciation

        return f"""当社{self.c.name}は、{self.c.established_date}の設立以来、{self.c.prefecture}を拠点として{self.c.industry}を営む企業である。主たる事業内容は{self.c.business_description}であり、現在、役員{self.c.officer_count}名、従業員{self.c.employee_count}名の体制で事業を運営している。

当社の経営を取り巻く環境は、近年大きく変化している。市場環境においては、{self.c.industry}に対する需要は堅調に推移しており、当社の売上高は2022年度{self.c.revenue_2022:,}円、2023年度{self.c.revenue_2023:,}円、2024年度{self.c.revenue_2024:,}円と着実に成長を遂げている。営業利益についても2022年度{self.c.operating_profit_2022:,}円、2023年度{self.c.operating_profit_2023:,}円、2024年度{self.c.operating_profit_2024:,}円と堅調に推移しており、当社の技術力と顧客からの信頼が数字として表れている。

しかしながら、事業成長を支える人材の確保については極めて厳しい状況に直面している。{self.c.industry}における有効求人倍率は{self.job_ratio}倍と高水準で推移しており、必要な人材を確保することが年々困難になっている。当社においても、{self.s.recruitment_period}にわたり継続的に求人活動を実施しているものの、{"応募者が極めて少なく" if self.s.applications == 0 else f"応募者数は{self.s.applications}名にとどまり"}、{"採用に至った人材は皆無であり" if self.s.hired == 0 else f"実際に採用に至ったのは{self.s.hired}名という"}厳しい結果となっている。

このような人手不足の状況下において、当社の競争力の源泉である技術力と品質を維持しながら、増加する顧客ニーズに対応していくためには、業務の省力化・効率化が不可欠な経営課題となっている。"""

    def generate_swot_analysis(self) -> str:
        """SWOT分析を生成"""
        return f"""【SWOT分析】

■強み（Strengths）
当社の最大の強みは、{self.c.established_date}の設立以来培ってきた{self.c.industry}における専門的な技術力とノウハウである。{self.c.business_description}に関する長年の経験に裏打ちされた高品質なサービス提供により、顧客からの厚い信頼を獲得している。また、役員{self.c.officer_count}名、従業員{self.c.employee_count}名という機動力のある組織体制により、顧客ニーズへの迅速な対応が可能である。

■弱み（Weaknesses）
一方で、{self.s.shortage_tasks}における業務効率の低さが課題である。従来型の手作業に依存した業務プロセスでは、1件あたりの作業時間が長く、増加する需要に十分対応できていない。また、慢性的な人手不足により、従業員への負担が過大となっている。

■機会（Opportunities）
省力化投資補助金を活用した{self.e.name}の導入は、当社にとって業務改革を実現する絶好の機会である。AI・デジタル技術の進展により、これまで自動化が困難であった業務も効率化が可能となっている。

■脅威（Threats）
{self.c.industry}における有効求人倍率は{self.job_ratio}倍と高水準で推移しており、人材確保の競争は今後さらに激化すると予想される。また、同業他社もデジタル化・省力化を進めており、対応が遅れれば競争力を失うリスクがある。"""

    def generate_section_1_2(self) -> str:
        """1-2 経営上の課題（PREP法、700字以上）"""
        return f"""当社が直面している最も深刻な経営課題は、慢性的な人手不足とそれに起因する従業員の過重労働である。

現在、{self.s.shortage_tasks}の業務を担当しているのは{self.s.current_workers}名であるが、業務量に対して適正な人員は{self.s.desired_workers}名が必要と考えている。すなわち、現状では{max(0, self.s.desired_workers - self.s.current_workers)}名の人員が不足している状態で業務を遂行せざるを得ない状況にある。

この人員不足を補うため、現場の従業員は月平均{self.s.overtime_hours}時間の残業を余儀なくされている。この数値は、厚生労働省が定める時間外労働の上限規制である月45時間に迫る水準であり、従業員の健康管理の観点からも早急な改善が求められている。長時間労働の常態化は、従業員の疲労蓄積による作業効率の低下を招くだけでなく、ミスや事故のリスクを高め、最悪の場合には貴重な人材の離職につながりかねない。

特に深刻なのは、{self.s.shortage_tasks}における作業負担である。この業務は従来、熟練した従業員の経験と勘に依存しており、1件あたり{self.l.current_hours}時間もの作業時間を要している。案件数の増加に伴い、この作業に費やす時間が増大し、他の重要業務に充てる時間が圧迫されている状況である。

さらに、新規人材の採用が困難な状況が続く中、既存従業員の高齢化も進行しており、技術やノウハウの継承という観点からも、早急に業務プロセスの見直しと省力化を図る必要性が高まっている。このまま対策を講じなければ、当社の事業継続そのものが危ぶまれる事態に陥りかねない。"""

    def generate_section_1_3(self) -> str:
        """1-3 動機・目的（PREP法、400字以上）"""
        # Phase 4: motivation_background を反映
        motivation_text = ""
        if self.data.motivation_background:
            motivation_text = f"\n\n本設備導入を決断した背景として、{self.data.motivation_background}という事情がある。"

        return f"""上記の経営課題を解決するため、当社は{self.e.name}の導入を決断した。{motivation_text}

本設備導入の最大の目的は、{self.s.shortage_tasks}における作業時間を大幅に削減し、従業員の過重労働を解消することにある。具体的には、現在1件あたり{self.l.current_hours}時間を要している作業を、本設備の導入により{self.l.target_hours}時間まで短縮することを目指している。これにより、作業時間を{self.l.reduction_rate:.0f}%削減し、月{self.s.overtime_hours}時間に及ぶ残業時間の大幅な圧縮を実現する。

省力化により創出された時間は、より付加価値の高い業務に充当する計画である。従業員が本来の専門性を発揮できる環境を整備することで、サービス品質の向上と顧客満足度の向上を図り、ひいては売上拡大と利益率の改善につなげていく。また、労働環境の改善は従業員の定着率向上にも寄与し、人材確保の面でもプラスの効果が期待できる。

本補助金を活用することで、当社の経営基盤を強化し、持続可能な成長を実現したい。"""

    def generate_section_2_1(self) -> str:
        """2-1 ビフォーアフター（PREP法、1000字以上）"""
        before_total = sum(p.time_minutes for p in self.data.before_processes)
        after_total = sum(p.time_minutes for p in self.data.after_processes)
        reduction_minutes = before_total - after_total

        text = f"""本事業において導入する{self.e.name}について、導入前後の業務プロセスの変化を詳細に説明する。

【導入前の業務プロセス】
現在、{self.s.shortage_tasks}の業務は、以下のプロセスで実施している。"""

        for p in self.data.before_processes:
            text += f"\n「{p.name}」工程では、{p.description}を行っており、所要時間は{p.time_minutes}分である。"

        text += f"""

これらの工程を合計すると、1サイクルあたり{before_total}分（約{before_total/60:.1f}時間）を要している。この作業を1日に複数回実施するため、{self.s.shortage_tasks}だけで1日あたり{self.l.current_hours}時間もの時間を費やしている状況である。作業の大部分は従業員の手作業に依存しており、膨大な資料との照合作業が必要となり、従業員の負担が極めて大きい。

【導入後の業務プロセス】
{self.e.name}を導入することで、業務プロセスは以下のように変化する。"""

        for p in self.data.after_processes:
            text += f"\n「{p.name}」工程は、{p.description}により{p.time_minutes}分で完了する。"

        # Phase 1: ゼロ除算防止
        reduction_pct = (reduction_minutes / before_total * 100) if before_total > 0 else 0

        text += f"""

導入後の合計所要時間は{after_total}分（約{after_total/60:.1f}時間）となる。導入前と比較して、{reduction_minutes}分（約{reduction_minutes/60:.1f}時間）の短縮、削減率にして{reduction_pct:.0f}%の省力化を実現する。

【工程別の省力化効果】
各工程における具体的な省力化効果は以下のとおりである。"""

        # 工程別の詳細分析
        process_pairs = list(zip(self.data.before_processes, self.data.after_processes))
        for bp, ap in process_pairs:
            saved = bp.time_minutes - ap.time_minutes
            if saved > 0:
                pct = saved / bp.time_minutes * 100 if bp.time_minutes > 0 else 0
                text += f"\n・「{bp.name}」工程：{bp.time_minutes}分→{ap.time_minutes}分（{saved}分削減、{pct:.0f}%減）。従来の{bp.description}を{ap.description}に置き換えることで効率化される。"
            else:
                text += f"\n・「{bp.name}」工程：{bp.time_minutes}分→{ap.time_minutes}分。本工程は人間の判断が必要であり、所要時間に変化はない。"

        # 最も効果の大きい工程を特定
        biggest = max(process_pairs, key=lambda pair: pair[0].time_minutes - pair[1].time_minutes)

        text += f"""

【省力化の仕組み】
{self.e.name}の主要機能として、{self.e.features if self.e.features else f"{self.e.name}による業務自動化・効率化機能"}が挙げられる。特に「{biggest[0].name}」工程においては、従来{biggest[0].description}に{biggest[0].time_minutes}分を要していたが、本設備の{biggest[1].description}機能により{biggest[1].time_minutes}分まで短縮される。これが本事業における最大の省力化ポイントである。

本設備の導入により、従業員は定型的・反復的な作業から解放され、顧客対応や品質管理といった人間の判断力が求められる高付加価値業務に集中できるようになる。1日あたりの削減時間は{self.l.reduction_hours:.1f}時間となり、月間では約{self.l.reduction_hours * Config.WORKING_DAYS_PER_MONTH:.0f}時間の業務時間を創出できる。"""

        return text

    def generate_section_2_2(self) -> str:
        """2-2 効果（PREP法、600字以上）"""
        # Phase 2: Config参照
        annual_saving = int(self.l.reduction_hours * Config.WORKING_DAYS_PER_MONTH * 12 * Config.HOURLY_WAGE)
        # Phase 4: time_utilization_plan を反映
        utilization_text = ""
        if self.data.time_utilization_plan:
            utilization_text = f"具体的には、{self.data.time_utilization_plan}に充てる計画である。"

        return f"""本事業の実施により期待される効果について、定量的・定性的の両面から説明する。

【定量的効果】
作業時間の削減効果として、1日あたり{self.l.reduction_hours:.1f}時間、月間では約{self.l.reduction_hours * Config.WORKING_DAYS_PER_MONTH:.0f}時間の業務時間を創出できる。この時間を人件費に換算すると、時給{Config.HOURLY_WAGE:,}円として年間約{annual_saving:,}円相当の効果となる。また、残業時間の削減により、割増賃金の支出も抑制される。現状の月{self.s.overtime_hours}時間の残業を半減できれば、年間で相当額の人件費削減が見込まれる。

【定性的効果】
まず、従業員の労働環境が大幅に改善される。長時間労働の解消により、従業員のワークライフバランスが向上し、心身の健康維持に寄与する。これは従業員の定着率向上につながり、採用難が続く現状において極めて重要な効果である。

次に、業務品質の安定化が期待できる。手作業に依存していた工程を自動化することで、ヒューマンエラーのリスクが大幅に低減される。一定の品質を安定して提供できることは、顧客からの信頼向上につながる。

さらに、創出された時間を活用して、より付加価値の高いサービスの提供や、新規顧客の開拓に注力することが可能となる。{utilization_text}これにより、売上の拡大と利益率の向上を実現し、持続的な事業成長の基盤を構築できる。"""

    def generate_section_3_1(self) -> str:
        """3-1 生産性向上（PREP法、700字以上）"""
        # Phase 2: Config参照
        base_added_value = self.c.operating_profit_2024 + int(self.c.revenue_2024 * Config.LABOR_COST_RATIO) + self.c.depreciation
        growth = Config.GROWTH_RATE

        # Phase 4: 賃上げ計画データの反映
        wage_detail = ""
        if self.data.wage_increase_rate > 0:
            wage_detail = f"当社は賃上げ率{self.data.wage_increase_rate}%を計画しており、"
            if self.data.wage_increase_target:
                wage_detail += f"対象は{self.data.wage_increase_target}、"
            if self.data.wage_increase_timing:
                wage_detail += f"{self.data.wage_increase_timing}より実施予定である。"
            else:
                wage_detail += "次年度より実施予定である。"

        growth_pct = (Config.GROWTH_RATE - 1) * 100
        salary_growth_pct = (Config.SALARY_GROWTH_RATE - 1) * 100

        return f"""本事業の実施により、当社は付加価値額の年率{growth_pct:.0f}%以上の向上を目指す。

【付加価値額の向上計画】
当社の付加価値額（営業利益＋人件費＋減価償却費）は、直近の2024年度実績で約{base_added_value:,}円である。本事業により省力化を実現し、業務効率を向上させることで、より多くの案件に対応可能となる。これにより、売上高の拡大を図りながら、付加価値額を毎年{growth_pct:.0f}%以上成長させていく計画である。

5年間の付加価値額推移の計画は以下のとおりである。
基準年度：約{base_added_value:,}円
1年目：約{int(base_added_value * growth):,}円（前年比+{growth_pct:.1f}%）
2年目：約{int(base_added_value * growth ** 2):,}円（前年比+{growth_pct:.1f}%）
3年目：約{int(base_added_value * growth ** 3):,}円（前年比+{growth_pct:.1f}%）
4年目：約{int(base_added_value * growth ** 4):,}円（前年比+{growth_pct:.1f}%）
5年目：約{int(base_added_value * growth ** 5):,}円（前年比+{growth_pct:.1f}%）

【給与支給総額の向上計画】
生産性向上により創出した利益の一部を原資として、従業員への還元を行う。具体的には、1人当たり給与支給総額の年平均成長率{salary_growth_pct:.1f}%以上を達成する計画である。{wage_detail}

【事業場内最低賃金の引上げ】
当社は、事業場内最低賃金について、{self.c.prefecture}の地域別最低賃金を30円以上上回る水準を維持することを表明する。

【投資回収計画】
本設備への投資額{self.f.total_investment:,}円は、省力化による人件費削減効果と売上拡大による利益増加により、約2〜3年で回収できる見込みである。"""


# =============================================================================
# 図解生成（Gemini API）
# =============================================================================

def generate_diagrams(data: HearingData, output_dir: str) -> Dict[str, str]:
    """全ての図解を生成（Phase 5: exponential backoff付きリトライ）"""
    if not GEMINI_AVAILABLE:
        print("  ⚠️ Gemini APIが利用できません")
        return {}

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        print("  ⚠️ GEMINI_API_KEY未設定")
        return {}

    print(f"\n🎨 図解を生成中（{Config.GEMINI_MODEL}）...")

    client = genai.Client(api_key=api_key)
    diagram_dir = Path(output_dir) / "diagrams"
    diagram_dir.mkdir(exist_ok=True)

    c, s, l, e, f = data.company, data.labor_shortage, data.labor_saving, data.equipment, data.funding
    diagrams = {}

    specs = [
        ("01_企業概要", f"企業概要図\n会社名:{c.name}\n業種:{c.industry}\n従業員:{c.employee_count}名\n設立:{c.established_date}\n事業:{c.business_description}"),
        ("02_SWOT分析", f"SWOT分析図（4象限）\n強み:専門技術、経験豊富\n弱み:人手不足、業務効率低下\n機会:省力化設備導入\n脅威:人材確保競争激化"),
        ("03_人手不足", f"人手不足状況図\n必要人員:{s.desired_workers}名\n現在:{s.current_workers}名\n不足:{s.desired_workers-s.current_workers}名\n残業:{s.overtime_hours}時間/月"),
        ("04_課題フロー", f"課題の連鎖図（矢印で連鎖を示す）\n業種:{c.industry}\n対象業務:{s.shortage_tasks}\n\n人手不足（現{s.current_workers}名/必要{s.desired_workers}名）→業務過多（{s.shortage_tasks}に1日{l.current_hours}時間）→残業増加（月{s.overtime_hours}時間）→品質低下・離職リスク→さらなる人手不足\n\n根本原因：手作業中心の業務プロセスが非効率"),
        ("05_設備概要", f"導入設備概要\n名称:{e.name}\n金額:{e.total_price:,}円\n特徴:AI活用、自動化"),
        ("06_ビフォーアフター", f"ビフォーアフター比較図（横棒グラフ形式で工程別に表示）\n設備名:{e.name}\n\n" + "\n".join([f"{bp.name}: 導入前{bp.time_minutes}分→導入後{ap.time_minutes}分" for bp, ap in zip(data.before_processes, data.after_processes)]) + f"\n\n合計: 導入前{l.current_hours}時間→導入後{l.target_hours}時間\n削減:{l.reduction_hours:.1f}時間（{l.reduction_rate:.0f}%削減）"),
        ("07_効果算定", f"省力化効果の定量分析図\n設備名:{e.name}\n\n削減時間:{l.reduction_hours:.1f}時間/日\n月間削減:{l.reduction_hours*22:.0f}時間\n年間削減:{l.reduction_hours*Config.WORKING_DAYS_PER_YEAR:.0f}時間\n削減率:{l.reduction_rate:.0f}%\n人件費換算:年間約{int(l.reduction_hours*Config.WORKING_DAYS_PER_MONTH*12*Config.HOURLY_WAGE):,}円相当"),
        ("12_業務フロー", f"現状の業務フロー図（フローチャート形式・左から右に工程を並べる）\n会社名:{c.name}\n業種:{c.industry}\n対象業務:{s.shortage_tasks}\n\n" + "→".join([f"{p.name}({p.time_minutes}分)" for p in data.before_processes]) + f"\n\n合計所要時間: {sum(p.time_minutes for p in data.before_processes)}分/サイクル\n問題点: 手作業中心で1日{l.current_hours}時間を要する"),
        ("13_工程別比較", f"工程別の省力化効果比較チャート（横棒グラフ：各工程の導入前vs導入後の所要時間を色分けで並べる）\n設備名:{e.name}\n\n" + "\n".join([f"{bp.name}: 導入前{bp.time_minutes}分→導入後{ap.time_minutes}分（{bp.time_minutes-ap.time_minutes}分削減）" for bp, ap in zip(data.before_processes, data.after_processes)]) + f"\n\n全体削減率: {l.reduction_rate:.0f}%"),
        ("08_実施体制", f"実施体制図\n代表者:{c.representative}\n責任者:{f.implementation_manager}\n従業員:{c.employee_count}名"),
        ("09_スケジュール", f"実施スケジュール\n1ヶ月目:契約発注\n2ヶ月目:納品設置\n3ヶ月目:試運転\n4ヶ月目:本格稼働"),
        ("10_5年計画", f"5年計画グラフ\n付加価値額:年率+{(Config.GROWTH_RATE-1)*100:.0f}%成長\n給与支給総額:年率+{(Config.SALARY_GROWTH_RATE-1)*100:.1f}%成長\n投資回収:約2-3年"),
        ("11_実施工程", f"""補助事業のスケジュール表（ガントチャート形式）を作成してください。

【表の構成】
- 縦軸：フェーズとタスク名
- 横軸：補助事業実施期間（3月～翌3月の13ヶ月）＋ 事業計画1～5年目

【フェーズとタスク】
0.構想設計: 事業目的・目標設定(3-5月)、課題・改善方針検討(3-6月)、事業計画作成(4-7月)、社内プロジェクト体制決定(4-6月)、投資採算性・投資規模決定(5-8月)、予算・調達計画策定(6-8月)
1.機能設計: システム要件定義(6-8月)、システム構成策定(7-9月)、機能一覧定義(8-10月)
2.周辺機器の手配: 機械装置発注(8-9月)、部品・原材料調達(8-11月)
3.機能試作・システム組み立て: システム設計(9-11月)、システム発注・開発(10-12月)
4.評価: テスト・リリース(11-12月)、課題・改善方針検討(12-1月)
5.調整改善: システム再設計(1-2月)
6.稼働・実装: セキュリティ対策(2-3月,1-2年目)、保守・管理(3月以降,1-5年目継続)

【デザイン】
- 青系統の配色
- 活動期間は矢印(⇨)で表示
- プロフェッショナルなビジネス文書スタイル
- 会社名:{c.name}
- 設備名:{e.name}"""),
    ]

    for diagram_id, prompt in specs:
        print(f"    📊 {diagram_id}...", end=" ")
        output_path = diagram_dir / f"{diagram_id}.png"

        # Phase 5: exponential backoff 付きリトライ
        success = False
        for attempt in range(Config.GEMINI_RETRY_MAX):
            try:
                response = client.models.generate_content(
                    model=Config.GEMINI_MODEL,
                    contents=f"以下の内容を示すビジネス図解を生成してください。日本語で、青系統の配色で、プロフェッショナルなスタイルで。\n\n{prompt}",
                    config=types.GenerateContentConfig(response_modalities=["IMAGE", "TEXT"])
                )

                for part in response.candidates[0].content.parts:
                    if hasattr(part, 'inline_data') and part.inline_data:
                        image_data = part.inline_data.data
                        if isinstance(image_data, str):
                            image_data = base64.b64decode(image_data)
                        with open(output_path, 'wb') as f_out:
                            f_out.write(image_data)
                        if os.path.getsize(output_path) > 1000:
                            diagrams[diagram_id] = str(output_path)
                            print("✅")
                            success = True
                            break
                if success:
                    break
                if attempt < Config.GEMINI_RETRY_MAX - 1:
                    delay = Config.GEMINI_RETRY_BASE_DELAY * (2 ** attempt)
                    print(f"⏳ リトライ({attempt + 2}/{Config.GEMINI_RETRY_MAX})...", end=" ")
                    time.sleep(delay)
            except Exception as ex:
                if attempt < Config.GEMINI_RETRY_MAX - 1:
                    delay = Config.GEMINI_RETRY_BASE_DELAY * (2 ** attempt)
                    print(f"⏳ エラー、リトライ({attempt + 2}/{Config.GEMINI_RETRY_MAX})...", end=" ")
                    time.sleep(delay)
                else:
                    print(f"❌ ({ex})")

        if not success and diagram_id not in diagrams:
            print("❌")

        time.sleep(Config.GEMINI_INTER_REQUEST_DELAY)

    return diagrams


# =============================================================================
# Word文書生成
# =============================================================================

def generate_business_plan_1_2(data: HearingData, diagrams: Dict[str, str], output_dir: str, template_path: Path):
    """事業計画書その1その2を生成"""
    print("\n📝 事業計画書（その1＋その2）を生成中...")

    output_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    shutil.copy(template_path, output_path)
    os.chmod(output_path, 0o644)

    doc = Document(output_path)
    gen = ContentGenerator(data)
    c, s, l, e, f = data.company, data.labor_shortage, data.labor_saving, data.equipment, data.funding

    manufacturer = e.manufacturer if e.manufacturer else "オーダーメイド開発"
    model_name = e.model if e.model else "カスタム仕様"

    # ヘルパー関数
    def get_unique_cells(row):
        unique, seen = [], set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen:
                seen.add(cid)
                unique.append(cell)
        return unique

    def clear_and_write(cell, text):
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""
        if cell.paragraphs:
            cell.paragraphs[0].text = text
        else:
            cell.text = text

    # ----- テーブル0: 事業者情報 -----
    print("    📋 事業者情報...")
    if len(doc.tables) > 0:
        t = doc.tables[0]
        info = [c.name, f"代表取締役  {c.representative}", f"{c.prefecture}{c.address}",
                c.industry, c.established_date, f"{c.officer_count}名 ／ {c.employee_count}名", c.url or ""]
        for i, val in enumerate(info):
            if i < len(t.rows) and len(t.rows[i].cells) > 1:
                t.rows[i].cells[1].text = val

    # ----- テーブル1: 事業計画名 -----
    print("    📋 事業計画名...")
    if len(doc.tables) > 1:
        doc.tables[1].rows[0].cells[0].text = f"{e.name}の導入による業務省力化と生産性向上"[:30]

    # ----- テーブル2: 概要 -----
    print("    📋 事業計画概要...")
    if len(doc.tables) > 2:
        doc.tables[2].rows[0].cells[0].text = f"当社は{c.industry}を営む企業である。{s.shortage_tasks}において人手不足が深刻であり、月{s.overtime_hours}時間の残業が発生している。{e.name}を導入し、作業時間を{l.reduction_rate:.0f}%削減することで、生産性向上と従業員の負担軽減を実現する。"

    # ----- テーブル3: 導入設備 -----
    print("    📋 導入設備情報...")
    if len(doc.tables) > 3:
        doc.tables[3].rows[0].cells[0].text = f"【設備名称】{e.name}\n【メーカー】{manufacturer}\n【型番】{model_name}\n【数量】{e.quantity}台\n【金額】{e.total_price:,}円（税抜）\n【購入先】{e.vendor}"

    # ----- テーブル4: ネストテーブル + 本文 -----
    if len(doc.tables) > 4:
        t4 = doc.tables[4]

        # ネストテーブル（事業者概要）
        print("    📋 事業者概要テーブル（ネスト）...")
        cell0 = t4.rows[0].cells[0]
        if cell0.tables:
            nested = cell0.tables[0]
            overview = gen.generate_business_overview_table_data()

            # 行0-5: テキスト項目
            text_items = ["事業者名", "経営理念", "経営戦略", "事業コンセプト", "事業内容", "長期的なビジョン"]
            for row_idx, key in enumerate(text_items):
                if row_idx < len(nested.rows):
                    uc = get_unique_cells(nested.rows[row_idx])
                    if len(uc) >= 2:
                        clear_and_write(uc[1], overview[key])

            # 行7-10: 直近実績
            fin_data = overview["直近実績"]
            fin_rows = [(7, "売上金額"), (8, "売上総利益"), (9, "営業利益"), (10, "従業員数")]
            for row_idx, key in fin_rows:
                if row_idx < len(nested.rows):
                    uc = get_unique_cells(nested.rows[row_idx])
                    if len(uc) >= 4:
                        vals = fin_data[key]
                        fmt = lambda v: f"{v:,}円" if key != "従業員数" else f"{v}名"
                        clear_and_write(uc[1], fmt(vals[0]))
                        clear_and_write(uc[2], fmt(vals[1]))
                        clear_and_write(uc[3], fmt(vals[2]))

        # 本文セクション
        print("    📋 本文セクション（PREP法）...")
        sections = {
            1: gen.generate_section_1_1() + "\n\n" + gen.generate_swot_analysis(),
            2: gen.generate_section_1_2(),
            3: gen.generate_section_1_3(),
            4: f"【導入設備の詳細】\n設備名称：{e.name}\nメーカー：{manufacturer}\n型番：{model_name}\n数量：{e.quantity}台\n金額：{e.total_price:,}円（税抜）\n購入先：{e.vendor}\nカタログ番号：{e.catalog_number or 'オーダーメイド'}\n\n【設備の特徴】\n{e.features}\n\n【投資金額の内訳】\n事業費総額：{f.total_investment:,}円\n補助金申請額：{f.subsidy_amount:,}円\n自己負担額：{f.self_funding:,}円",
            5: gen.generate_section_2_1(),
            6: gen.generate_section_2_2(),
            8: gen.generate_section_3_1(),
            9: f"【資金調達計画】\n事業費総額：{f.total_investment:,}円\nうち補助金：{f.subsidy_amount:,}円\nうち自己資金：{f.self_funding:,}円\n\n自己資金については、当社の内部留保および取引銀行である{f.bank_name}からの借入により調達する予定である。\n\n【投資回収計画】\n本設備への投資は、省力化による人件費削減効果と売上拡大による利益増加により、約2〜3年での回収を見込んでいる。",
            10: f"【実施体制】\n統括責任者：{c.representative}（代表取締役）\n実施責任者：{f.implementation_manager}\n従業員{c.employee_count}名と連携して実施\n\n【スケジュール】\n実施期間：{f.implementation_period}\n\n1ヶ月目：契約・発注\n2ヶ月目：設備納品・設置工事\n3ヶ月目：試運転・調整・従業員教育\n4ヶ月目以降：本格稼働・効果測定",
            11: f"【人手不足の状況】\n当社は「限られた人手で業務を遂行するため、直近の従業員の平均残業時間が30時間を超えている」状況に該当する。直近12ヶ月の平均残業時間：月{s.overtime_hours}時間\n\n【オーダーメイド性】\n本設備は当社の業務に特化したカスタマイズを施す。{e.features}\n\n【賃上げ計画の表明】\n・1人当たり給与支給総額の年平均成長率：{(Config.SALARY_GROWTH_RATE - 1) * 100:.1f}%以上\n・事業場内最低賃金：{c.prefecture}の地域別最低賃金を30円以上上回る水準"
        }

        for row_idx, content in sections.items():
            if row_idx < len(t4.rows):
                cell = t4.rows[row_idx].cells[0]
                existing = cell.text
                cell.text = existing.rstrip() + "\n\n" + content.strip()
                print(f"      ✅ セクション{row_idx}（{len(content)}文字）")

        # 図解挿入
        if diagrams:
            print("    🖼️ 図解挿入...")
            mapping = {1: ["01_企業概要", "02_SWOT分析"], 2: ["03_人手不足", "04_課題フロー", "12_業務フロー"],
                       4: ["05_設備概要"], 5: ["06_ビフォーアフター", "13_工程別比較"], 6: ["07_効果算定"],
                       10: ["08_実施体制", "09_スケジュール"], 8: ["10_5年計画"]}
            for row_idx, ids in mapping.items():
                if row_idx < len(t4.rows):
                    cell = t4.rows[row_idx].cells[0]
                    for did in ids:
                        if did in diagrams:
                            try:
                                para = cell.add_paragraph()
                                para.add_run().add_picture(diagrams[did], width=Inches(5.5))
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as ex:
                                print(f"      ⚠️ 図解挿入エラー ({did}): {ex}")

    # スケジュール図を追加（11_実施工程）
    if "11_実施工程" in diagrams:
        print("    📅 補助事業スケジュール図...")
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run("補助事業のスケジュール（参考）")
        run.bold = True
        run.font.size = Pt(14)
        p = doc.add_paragraph()
        try:
            p.add_run().add_picture(diagrams["11_実施工程"], width=Inches(6.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as ex:
            print(f"      ⚠️ スケジュール図挿入エラー: {ex}")

    doc.save(output_path)
    print(f"  ✅ 保存完了: {output_path}")


def add_schedule_table(doc, data: HearingData):
    """補助事業のスケジュール表をWord表形式で追加"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import RGBColor

    base_year = 2026  # 交付決定想定年度

    # ページ区切り
    doc.add_page_break()

    # タイトル
    p = doc.add_paragraph()
    run = p.add_run("補助事業のスケジュール（参考）")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run(
        "前述の補助事業の内容に沿い機械装置等の取得時期や技術の導入時期を含めた"
        "スケジュールを示してください。記載例ですので適宜使いやすいように作成してください。"
    )
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run("≪事業計画期間にかかるスケジュール≫")
    run.bold = True
    run.font.size = Pt(11)

    # フェーズ定義: (フェーズ名, [(タスク名, 活動月リスト, 活動年リスト)])
    # 月: 0=3月, 1=4月, ..., 9=12月, 10=1月, 11=2月, 12=3月(翌年)
    # 年: 0=1年目, ..., 4=5年目
    phases = [
        ("0．構想設計", [
            ("事業目的・目標設定", list(range(0, 3)), []),
            ("課題・改善方針検討", list(range(0, 4)), []),
            ("事業計画作成", list(range(1, 5)), []),
            ("社内プロジェクト体\n制決定", list(range(1, 4)), []),
            ("投資採算性・投資規\n模決定", list(range(2, 6)), []),
            ("予算・調達計画策定", list(range(3, 6)), []),
        ]),
        ("1．機能設計", [
            ("システム要件定義", list(range(3, 6)), []),
            ("システム構成策定", list(range(4, 7)), []),
            ("機能一覧定義", list(range(5, 8)), []),
        ]),
        ("2．周辺機器の手配", [
            ("機械装置発注", list(range(5, 7)), []),
            ("部品・原材料調達", list(range(5, 9)), []),
        ]),
        ("3．機能試作、シス\nテム組み立て", [
            ("システム設計", list(range(6, 9)), []),
            ("システム発注・開発", list(range(7, 10)), []),
        ]),
        ("4．評価", [
            ("テスト・リリース", list(range(8, 10)), []),
            ("課題・改善方針検討", list(range(9, 11)), []),
        ]),
        ("5．調整改善", [
            ("システム再設計", list(range(10, 12)), []),
        ]),
        ("6．稼働・実装", [
            ("セキュリティ対策", list(range(11, 13)), [0, 1]),
            ("保守・管理", [12], [0, 1, 2, 3, 4]),
        ]),
    ]

    total_tasks = sum(len(tasks) for _, tasks in phases)
    HEADER_ROWS = 3
    TOTAL_COLS = 20  # 2(フェーズ+タスク) + 13(月) + 5(年)

    table = doc.add_table(rows=HEADER_ROWS + total_tasks, cols=TOTAL_COLS)
    table.style = 'Table Grid'

    def shade_cell(cell, color="B4C6E7"):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell(cell, text, size=7, bold=False, align=None, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        if align:
            p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def mark_active(cell, color="B4C6E7"):
        """活動期間セルにマーカーと背景色を設定"""
        shade_cell(cell, color)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("⇨")
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0x30, 0x60, 0xA0)

    # --- ヘッダー行0: 大見出し ---
    table.cell(0, 0).merge(table.cell(0, 1))
    set_cell(table.cell(0, 0), "", 7)
    table.cell(0, 2).merge(table.cell(0, 14))
    set_cell(table.cell(0, 2), "補助事業実施期間", 8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    table.cell(0, 15).merge(table.cell(0, 19))
    set_cell(table.cell(0, 15), "事業計画期間", 8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- ヘッダー行1: 年度ラベル ---
    table.cell(1, 0).merge(table.cell(1, 1))
    set_cell(table.cell(1, 0), "", 7)
    table.cell(1, 2).merge(table.cell(1, 14))
    set_cell(table.cell(1, 2), "", 7)
    for i in range(5):
        set_cell(table.cell(1, 15 + i), f"事業計画{i+1}年目", 6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- ヘッダー行2: 月＋年度期間 ---
    set_cell(table.cell(2, 0), "", 7)
    set_cell(table.cell(2, 1), "", 7)
    month_labels = ["3\n月", "4\n月", "5\n月", "6\n月", "7\n月", "8\n月",
                    "9\n月", "10\n月", "11\n月", "12\n月", "1", "2", "3"]
    for i, label in enumerate(month_labels):
        set_cell(table.cell(2, 2 + i), label, 6, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(5):
        ys = base_year + i
        ye = ys + 1
        if i == 0:
            label = f"※{ys}年4月～\n{ye}年3月"
        else:
            label = f"{ys}年4月～{ye}\n年3月"
        set_cell(table.cell(2, 15 + i), label, 5, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- データ行 ---
    current_row = HEADER_ROWS
    for phase_name, tasks in phases:
        start_row = current_row
        for task_name, active_months, active_years in tasks:
            row = current_row
            set_cell(table.cell(row, 1), task_name, 7)
            for m in active_months:
                if 0 <= m <= 12:
                    mark_active(table.cell(row, 2 + m))
            for y in active_years:
                if 0 <= y <= 4:
                    mark_active(table.cell(row, 15 + y))
            current_row += 1

        end_row = current_row - 1
        if start_row < end_row:
            table.cell(start_row, 0).merge(table.cell(end_row, 0))
        set_cell(table.cell(start_row, 0), phase_name, 7)

    # フッター注記
    p = doc.add_paragraph()
    run = p.add_run("事業計画は事業者ごとの決算期")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def generate_business_plan_3(data: HearingData, output_dir: str, template_path: Path):
    """
    事業計画書その3を生成（openpyxl方式）
    """
    print("\n📊 事業計画書（その3）を生成中...")

    c, e, f = data.company, data.equipment, data.funding

    # 出力先にコピー
    output_path = Path(output_dir) / "事業計画書_その3_完成版.xlsx"
    shutil.copy(template_path, output_path)
    os.chmod(output_path, 0o644)

    try:
        wb = openpyxl.load_workbook(output_path)

        # Phase 6: _find_sheet_in_workbook でシート検索
        # === 別紙1（省力化計算シート）===
        ws1 = _find_sheet_in_workbook(wb, ["別紙1", "省力化"])
        if ws1:
            print(f"    📋 別紙1: {ws1.title}")

            # 導入前工程（C11〜C16, D11〜D16）
            for i, p in enumerate(data.before_processes[:6]):
                ws1[f'C{11+i}'] = p.name
                ws1[f'D{11+i}'] = p.time_minutes

            # 導入後工程（I11〜I16, J11〜J16）
            for i, p in enumerate(data.after_processes[:6]):
                ws1[f'I{11+i}'] = p.name
                ws1[f'J{11+i}'] = p.time_minutes

            print("      ✅ 工程データ入力完了")

        # === 別紙3（投資回収期間計算シート）===
        ws3 = _find_sheet_in_workbook(wb, ["別紙3", "投資回収"])
        if ws3:
            print(f"    📋 別紙3: {ws3.title}")

            # Phase 2: Config参照
            ws3['C6'] = f.total_investment  # 投資総額
            ws3['H6'] = Config.WORKING_DAYS_PER_YEAR  # 年間稼働日数
            ws3['J6'] = Config.HOURLY_WAGE  # 人件費単価
            ws3['L6'] = int(e.total_price / Config.DEPRECIATION_YEARS)  # 減価償却費

            print("      ✅ 投資回収データ入力完了")

        # === 参考書式（事業計画目標値算出シート）===
        ws_ref = _find_sheet_in_workbook(wb, ["参考書式", "目標値"])
        if ws_ref:
            print(f"    📋 参考書式: {ws_ref.title}")

            # --- ラベル行を動的検索 ---
            def find_row_by_label(ws, keywords, search_cols=('A', 'B', 'C', 'D'), max_row=60):
                """シート内でキーワードに一致するラベル行を見つける"""
                for row in range(1, max_row + 1):
                    for col_letter in search_cols:
                        val = ws[f'{col_letter}{row}'].value
                        if val:
                            for kw in keywords:
                                if kw in str(val):
                                    return row
                return None

            # 行番号を検索（フォールバック付き）
            row_revenue = find_row_by_label(ws_ref, ["売上高"]) or 26
            row_operating_profit = find_row_by_label(ws_ref, ["営業利益"])
            row_labor_cost = find_row_by_label(ws_ref, ["人件費"])
            row_depreciation = find_row_by_label(ws_ref, ["減価償却費"])
            row_added_value = find_row_by_label(ws_ref, ["付加価値額"])
            row_officers = find_row_by_label(ws_ref, ["役員数"]) or 37
            row_employees = find_row_by_label(ws_ref, ["従業員数"]) or 38
            row_salary_total = find_row_by_label(ws_ref, ["給与支給総額"]) or 44
            row_salary_employees = find_row_by_label(ws_ref, ["給与対象"]) or 45

            base_revenue = c.revenue_2024
            base_op_profit = c.operating_profit_2024

            # 人件費: 決算書PDFから取得 or 売上高×労働分配率で推計
            base_labor_cost = c.labor_cost if c.labor_cost > 0 else int(base_revenue * Config.LABOR_COST_RATIO)
            # 減価償却費: 決算書PDFから取得 or 設備投資額÷耐用年数で推計
            base_depreciation = c.depreciation if c.depreciation > 0 else int(e.total_price / Config.DEPRECIATION_YEARS)
            # 給与支給総額: 決算書PDFから取得 or 売上高×給与比率で推計
            base_salary = c.total_salary if c.total_salary > 0 else int(base_revenue * Config.SALARY_RATIO)
            # 付加価値額 = 営業利益 + 人件費 + 減価償却費
            base_added_value = base_op_profit + base_labor_cost + base_depreciation

            # E列=基準, G〜K列=1〜5年目
            cols = ['E', 'G', 'H', 'I', 'J', 'K']

            for i, col in enumerate(cols):
                growth = Config.GROWTH_RATE ** i
                salary_growth = Config.SALARY_GROWTH_RATE ** i

                # 売上高
                ws_ref[f'{col}{row_revenue}'] = int(base_revenue * growth)
                # 営業利益
                if row_operating_profit:
                    ws_ref[f'{col}{row_operating_profit}'] = int(base_op_profit * growth)
                # 人件費
                if row_labor_cost:
                    ws_ref[f'{col}{row_labor_cost}'] = int(base_labor_cost * salary_growth)
                # 減価償却費
                if row_depreciation:
                    ws_ref[f'{col}{row_depreciation}'] = int(base_depreciation)
                # 付加価値額
                if row_added_value:
                    av_op = int(base_op_profit * growth)
                    av_lc = int(base_labor_cost * salary_growth)
                    av_dep = int(base_depreciation)
                    ws_ref[f'{col}{row_added_value}'] = av_op + av_lc + av_dep
                # 役員数
                ws_ref[f'{col}{row_officers}'] = c.officer_count
                # 従業員数
                ws_ref[f'{col}{row_employees}'] = c.employee_count
                # 給与支給総額（年率2.5%成長）
                ws_ref[f'{col}{row_salary_total}'] = int(base_salary * salary_growth)
                # 給与対象従業員数
                ws_ref[f'{col}{row_salary_employees}'] = c.employee_count

            # 成長率の確認ログ
            year5_added_value = int(base_op_profit * Config.GROWTH_RATE**5) + int(base_labor_cost * Config.SALARY_GROWTH_RATE**5) + int(base_depreciation)
            if base_added_value > 0:
                av_annual_growth = ((year5_added_value / base_added_value) ** (1/5) - 1) * 100
                print(f"      📊 付加価値額: 基準{base_added_value:,}円 → 5年目{year5_added_value:,}円（年率{av_annual_growth:.1f}%）")
            year5_salary = int(base_salary * Config.SALARY_GROWTH_RATE**5)
            if base_salary > 0:
                sal_annual_growth = ((year5_salary / base_salary) ** (1/5) - 1) * 100
                print(f"      📊 給与支給総額: 基準{base_salary:,}円 → 5年目{year5_salary:,}円（年率{sal_annual_growth:.1f}%）")

            print("      ✅ 目標値データ入力完了")

        # 保存
        wb.save(output_path)
        wb.close()

        print(f"  ✅ 保存完了: {output_path.name}")
        print("    ⚠️ 注意: ファイルが開けない場合は手動でコピーが必要です")

    except Exception as ex:
        print(f"    ⚠️ openpyxlエラー: {ex}")
        print("    📝 2ファイル方式にフォールバック...")
        # フォールバック: 2ファイル方式
        from openpyxl import Workbook
        data_file = Path(output_dir) / "事業計画書_その3_入力データ.xlsx"
        wb_new = Workbook()
        ws1 = wb_new.active
        ws1.title = "別紙1_工程データ"
        for i, p in enumerate(data.before_processes):
            ws1.cell(row=2+i, column=1, value=p.name)
            ws1.cell(row=2+i, column=2, value=p.time_minutes)
        for i, p in enumerate(data.after_processes):
            ws1.cell(row=2+i, column=4, value=p.name)
            ws1.cell(row=2+i, column=5, value=p.time_minutes)
        wb_new.save(data_file)
        print(f"    ✅ 入力データ: {data_file.name}")


def generate_other_documents(data: HearingData, output_dir: str, template_dir: Path):
    """その他の書類を生成（openpyxlでデータ入力）"""
    print("\n📄 その他の書類を生成中...")

    from openpyxl.cell.cell import MergedCell

    c = data.company
    f = data.funding

    def safe_write(ws, cell_addr, value):
        """マージセルでも安全に書き込む"""
        cell = ws[cell_addr]
        if isinstance(cell, MergedCell):
            for mc in ws.merged_cells.ranges:
                if cell.coordinate in mc:
                    ws.cell(mc.min_row, mc.min_col).value = value
                    return
        else:
            cell.value = value

    # === 1. 役員名簿 ===
    try:
        src = template_dir / "役員名簿_様式.xlsx"
        dst = Path(output_dir) / "役員名簿_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        # Phase 6: ファジーマッチング
        ws = _find_sheet_in_workbook(wb, ["役員名簿", "役員"]) or wb[wb.sheetnames[0]]
        safe_write(ws, 'D5', c.name)  # 法人名
        safe_write(ws, 'D7', c.officer_count)  # 役員数
        safe_write(ws, 'D8', 0)  # 大企業所属人数
        safe_write(ws, 'D9', 0)  # みなし大企業所属人数
        for i, off in enumerate(data.officers[:10]):
            row = 15 + i
            ws[f'B{row}'] = i + 1
            ws[f'C{row}'] = off.position
            # Phase 6: _split_name で全角スペース対応
            last_name, first_name = _split_name(off.name)
            ws[f'D{row}'] = last_name
            ws[f'E{row}'] = first_name
        wb.save(dst)
        wb.close()
        print(f"    ✅ 役員名簿_完成版.xlsx（{c.officer_count}名）")
    except Exception as e:
        print(f"    ⚠️ 役員名簿エラー: {e}")

    # === 2. 従業員名簿 ===
    try:
        src = template_dir / "従業員名簿_様式.xlsx"
        dst = Path(output_dir) / "従業員名簿_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        # Phase 6: ファジーマッチング
        ws = _find_sheet_in_workbook(wb, ["労働者名簿", "従業員名簿", "従業員"]) or wb[wb.sheetnames[0]]
        safe_write(ws, 'C5', c.name)
        safe_write(ws, 'C7', c.employee_count)
        for i, emp in enumerate(data.employees[:50]):
            row = 12 + i
            ws[f'B{row}'] = i + 1
            # Phase 6: _split_name で全角スペース対応
            last_name, first_name = _split_name(emp.name)
            ws[f'C{row}'] = last_name
            ws[f'D{row}'] = first_name
            if emp.birth_date:
                ws[f'E{row}'] = emp.birth_date
        wb.save(dst)
        wb.close()
        print(f"    ✅ 従業員名簿_完成版.xlsx（{c.employee_count}名）")
    except Exception as e:
        print(f"    ⚠️ 従業員名簿エラー: {e}")

    # === 3. 株主・出資者名簿 ===
    try:
        src = template_dir / "株主出資者名簿_様式.xlsx"
        dst = Path(output_dir) / "株主出資者名簿_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        # Phase 6: ファジーマッチング
        ws = _find_sheet_in_workbook(wb, ["株主名簿", "株主"]) or wb[wb.sheetnames[0]]
        safe_write(ws, 'C5', c.name)
        safe_write(ws, 'C6', c.capital)
        for i, sh in enumerate(data.shareholders[:20]):
            row = 14 + i
            ws[f'B{row}'] = i + 1
            ws[f'C{row}'] = sh.name
            ws[f'D{row}'] = sh.shares
        wb.save(dst)
        wb.close()
        print(f"    ✅ 株主出資者名簿_完成版.xlsx（{len(data.shareholders)}名）")
    except Exception as e:
        print(f"    ⚠️ 株主名簿エラー: {e}")

    # === 4. 事業実施場所リスト ===
    try:
        src = template_dir / "事業実施場所リスト_様式.xlsx"
        dst = Path(output_dir) / "事業実施場所リスト_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        # Phase 6: ファジーマッチング
        ws = _find_sheet_in_workbook(wb, ["所在地リスト", "所在地", "事業実施場所"]) or wb[wb.sheetnames[0]]
        ws['C20'] = c.postal_code.replace('-', '') if c.postal_code else ''
        ws['C21'] = c.prefecture
        addr_parts = c.address.replace(c.prefecture, '').strip() if c.address else ''
        ws['C22'] = addr_parts[:10] if addr_parts else ''
        ws['C23'] = addr_parts[10:20] if len(addr_parts) > 10 else ''
        ws['C24'] = addr_parts[20:] if len(addr_parts) > 20 else ''
        ws['C26'] = c.name + ' 本社'
        ws['C27'] = c.phone
        wb.save(dst)
        wb.close()
        print(f"    ✅ 事業実施場所リスト_完成版.xlsx")
    except Exception as e:
        print(f"    ⚠️ 事業実施場所リストエラー: {e}")

    # === 5. 他の補助金使用実績 ===
    try:
        src = template_dir / "他の補助金使用実績_様式.xlsx"
        dst = Path(output_dir) / "他の補助金使用実績_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        # Phase 6: ファジーマッチング
        ws = _find_sheet_in_workbook(wb, ["Sheet1", "他の補助金", "使用実績"]) or wb[wb.sheetnames[0]]
        ws['C25'] = "なし"
        ws['C26'] = "-"
        ws['C27'] = "-"
        ws['C28'] = "-"
        wb.save(dst)
        wb.close()
        print(f"    ✅ 他の補助金使用実績_完成版.xlsx")
    except Exception as e:
        print(f"    ⚠️ 他の補助金使用実績エラー: {e}")

    # === 6. 給与支給総額確認書 ===
    try:
        src = template_dir / "給与支給総額確認書_様式.xlsx"
        dst = Path(output_dir) / "給与支給総額確認書_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        if "宣誓書" in wb.sheetnames:
            ws = wb["宣誓書"]
            safe_write(ws, 'C8', c.name)
            safe_write(ws, 'E8', c.representative)
        for sname in wb.sheetnames:
            if "直近決算" in sname and "記入例" not in sname and "未満" not in sname:
                ws = wb[sname]
                safe_write(ws, 'C5', c.name)
                base_salary = c.total_salary if c.total_salary > 0 else int(c.revenue_2024 * Config.SALARY_RATIO)
                safe_write(ws, 'E10', base_salary)
                safe_write(ws, 'E11', c.employee_count)
                break
        wb.save(dst)
        wb.close()
        print(f"    ✅ 給与支給総額確認書_完成版.xlsx")
    except Exception as e:
        print(f"    ⚠️ 給与支給総額確認書エラー: {e}")

    # === 7. 賃金引上げ要件（事業場内） ===
    try:
        src = template_dir / "賃金引上げ要件_事業場内_様式.xlsx"
        dst = Path(output_dir) / "賃金引上げ要件_事業場内_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        if "確認書" in wb.sheetnames:
            ws = wb["確認書"]
            safe_write(ws, 'C7', c.name)
            safe_write(ws, 'E7', c.representative)
        wb.save(dst)
        wb.close()
        print(f"    ✅ 賃金引上げ要件_事業場内_完成版.xlsx")
    except Exception as e:
        print(f"    ⚠️ 賃金引上げ要件_事業場内エラー: {e}")

    # === 8. 賃金引上げ要件（地域別） ===
    try:
        src = template_dir / "賃金引上げ要件_地域別_様式.xlsx"
        dst = Path(output_dir) / "賃金引上げ要件_地域別_完成版.xlsx"
        shutil.copy(src, dst)
        wb = openpyxl.load_workbook(dst)
        if "確認書" in wb.sheetnames:
            ws = wb["確認書"]
            safe_write(ws, 'C7', c.name)
            safe_write(ws, 'E7', c.representative)
        wb.save(dst)
        wb.close()
        print(f"    ✅ 賃金引上げ要件_地域別_完成版.xlsx")
    except Exception as e:
        print(f"    ⚠️ 賃金引上げ要件_地域別エラー: {e}")

    # === 9. 金融機関確認書（Word） ===
    try:
        src = template_dir / "金融機関確認書_様式.docx"
        dst = Path(output_dir) / "金融機関確認書_完成版.docx"
        shutil.copy(src, dst)
        doc = Document(str(dst))
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    txt = cell.text
                    if '法人名' in txt or '申請者名' in txt or '事業者名' in txt:
                        if i + 1 < len(row.cells):
                            row.cells[i + 1].text = c.name
                    if '代表者名' in txt:
                        if i + 1 < len(row.cells):
                            row.cells[i + 1].text = c.representative
                    if '金融機関名' in txt:
                        if i + 1 < len(row.cells):
                            row.cells[i + 1].text = f.bank_name
        doc.save(str(dst))
        print(f"    ✅ 金融機関確認書_完成版.docx")
    except Exception as e:
        print(f"    ⚠️ 金融機関確認書エラー: {e}")


# =============================================================================
# 書類生成（1回分）
# =============================================================================

def _run_generation(data: HearingData, output_dir: str, template_dir, diagrams: dict):
    """書類一式を生成する（1回分の実行）"""
    template_dir = Path(template_dir)
    Path(output_dir).mkdir(exist_ok=True, parents=True)

    t = template_dir / "事業計画書_その1その2_様式.docx"
    if t.exists():
        generate_business_plan_1_2(data, diagrams, str(output_dir), t)

    t = template_dir / "事業計画書_その3_様式.xlsx"
    if t.exists():
        generate_business_plan_3(data, str(output_dir), t)

    generate_other_documents(data, str(output_dir), template_dir)


def _fix_text_holes_in_docx(output_dir: str, data: HearingData) -> list:
    """docx内のテキスト穴あき（プレースホルダー空白）を修正する"""
    docx_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return []

    doc = Document(str(docx_path))
    fixes = []

    # 修正マッピング: パターン → 置換テキスト生成関数
    e = data.equipment
    features = e.features if e.features else f"{e.name}による業務自動化・効率化機能"

    replacements = {
        "主要機能として、が挙げられる": f"主要機能として、{features}が挙げられる",
    }

    # 未丸め小数値の修正パターン
    decimal_pattern = re.compile(r"(\d+)\.(\d{6,})(\s*時間|\s*分)")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    text = original

                    # 固定パターン置換
                    for old_text, new_text in replacements.items():
                        if old_text in text:
                            text = text.replace(old_text, new_text)
                            fixes.append(f"穴あき修正: '{old_text}' → '{new_text[:40]}...'")

                    # 未丸め小数値修正
                    def _round_decimal(m):
                        whole = m.group(1)
                        decimal = m.group(2)
                        unit = m.group(3)
                        rounded = round(float(f"{whole}.{decimal}"), 1)
                        return f"{rounded}{unit}"
                    text = decimal_pattern.sub(_round_decimal, text)
                    if text != original:
                        if "小数丸め" not in str(fixes):
                            fixes.append("小数値丸め: 未整形の小数値を修正")

                    if text != original:
                        para.text = text

    for para in doc.paragraphs:
        original = para.text
        text = original
        for old_text, new_text in replacements.items():
            if old_text in text:
                text = text.replace(old_text, new_text)
                if f"穴あき修正: '{old_text}'" not in str(fixes):
                    fixes.append(f"穴あき修正: '{old_text}' → '{new_text[:40]}...'")
        text = decimal_pattern.sub(_round_decimal, text)
        if text != original:
            para.text = text

    if fixes:
        doc.save(str(docx_path))

    return fixes


def _apply_fixes(issues: list, data: HearingData) -> list:
    """スコアリング結果のissuesを解析し、パラメータを自動修正する。
    適用した修正のリストを返す。"""
    fixes_applied = []

    for issue in issues:
        action = issue.get("action", "")

        if action == "increase_growth_rate":
            old = Config.GROWTH_RATE
            Config.GROWTH_RATE = min(Config.GROWTH_RATE + 0.005, 1.10)  # 上限10%
            if Config.GROWTH_RATE != old:
                fixes_applied.append(f"GROWTH_RATE: {old} -> {Config.GROWTH_RATE}")

        elif action == "increase_salary_rate":
            old = Config.SALARY_GROWTH_RATE
            Config.SALARY_GROWTH_RATE = min(Config.SALARY_GROWTH_RATE + 0.005, 1.05)  # 上限5%
            if Config.SALARY_GROWTH_RATE != old:
                fixes_applied.append(f"SALARY_GROWTH_RATE: {old} -> {Config.SALARY_GROWTH_RATE}")

        elif action == "increase_text" or action == "increase_section_text":
            # テキスト不足はテンプレートで対応済みのため、再生成で解決を試みる
            if "テキスト再生成" not in [f.split(":")[0] for f in fixes_applied]:
                fixes_applied.append("テキスト再生成: リトライ")

        elif action == "fix_text_holes":
            # テキスト穴あき修正はdocx直接編集で対応
            if "テキスト穴あき修正" not in [f.split(":")[0] for f in fixes_applied]:
                fixes_applied.append("テキスト穴あき修正: docx直接編集")

        elif action == "fix_value_inconsistency":
            # 書類間整合性は再生成で対応（Config値が統一されていれば解決）
            if "数値整合性修正" not in [f.split(":")[0] for f in fixes_applied]:
                fixes_applied.append("数値整合性修正: 再生成で統一")

        elif action == "fix_negative_profit":
            # 営業利益マイナスは成長率増加で対応
            old = Config.GROWTH_RATE
            Config.GROWTH_RATE = min(Config.GROWTH_RATE + 0.01, 1.10)
            if Config.GROWTH_RATE != old:
                fixes_applied.append(f"営業利益修正: GROWTH_RATE {old} -> {Config.GROWTH_RATE}")

    return fixes_applied


def _extract_docx_text(output_dir: str) -> str:
    """事業計画書docxから全テキストを抽出する"""
    docx_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return ""
    doc = Document(str(docx_path))
    texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        texts.append(t)
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            texts.append(t)
    return "\n\n".join(texts)


def _write_text_to_docx(output_dir: str, rewritten_text: str):
    """リライト済みテキストを事業計画書docxのテーブルセルに書き戻す"""
    docx_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return

    doc = Document(str(docx_path))

    # セクション番号→リライト済みテキストのマッピングを構築
    # リライト済みテキストをセクションヘッダーで分割
    import re
    section_map = {}
    current_key = None
    current_lines = []

    for line in rewritten_text.split("\n"):
        # セクションヘッダー検出（【...】パターン）
        header_match = re.match(r"^【(.+?)】", line.strip())
        if header_match:
            if current_key and current_lines:
                section_map[current_key] = "\n".join(current_lines).strip()
            current_key = header_match.group(1)
            current_lines = [line]
        elif current_key:
            current_lines.append(line)

    if current_key and current_lines:
        section_map[current_key] = "\n".join(current_lines).strip()

    if not section_map:
        # セクション分割できない場合、全体を最大のテーブルセルに書き込む
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if len(cell.text) > 500:
                        cell.text = rewritten_text
                        doc.save(str(docx_path))
                        return
        return

    # テーブルセルをスキャンし、対応するセクションのテキストを置換
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                for key, new_text in section_map.items():
                    if key in cell_text and len(cell_text) > 200:
                        cell.text = new_text
                        break

    doc.save(str(docx_path))


def _run_deai_phase(
    output_dir: str,
    industry: str,
    target_ai_score: int = 85,
    max_rounds: int = 3,
    on_progress=None,
) -> dict:
    """AI臭除去フェーズ: docxテキスト抽出→スコアリング→リライト→書き戻し

    Returns:
        dict: {ai_score, ai_rounds, ai_history, skipped}
    """
    # ai_smell_score をインポート
    skill_scripts = Path.home() / ".claude" / "skills" / "shoryokuka-review-deai" / "scripts"
    if not skill_scripts.exists():
        print("  AI臭除去スキルが未インストール。スキップします。")
        return {"ai_score": None, "ai_rounds": 0, "ai_history": [], "skipped": True}

    import importlib.util
    spec = importlib.util.spec_from_file_location("ai_smell_score", str(skill_scripts / "ai_smell_score.py"))
    ai_smell = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(ai_smell)

    # テキスト抽出
    text = _extract_docx_text(output_dir)
    if not text or len(text) < 100:
        print("  事業計画書テキストが短すぎます。AI臭除去をスキップ。")
        return {"ai_score": None, "ai_rounds": 0, "ai_history": [], "skipped": True}

    # 初回スコアリング
    result = ai_smell.calculate_score(text)
    ai_score = result["total_score"]
    ai_history = [{"round": 0, "score": ai_score, "grade": result["grade"]}]
    print(f"\n  AI臭スコア（初回）: {ai_score}/100 ({result['grade']})")

    if on_progress:
        on_progress("ai_smell_initial", ai_score, result)

    if ai_score >= target_ai_score:
        print(f"  AI臭スコア {ai_score} >= {target_ai_score}。リライト不要。")
        return {"ai_score": ai_score, "ai_rounds": 0, "ai_history": ai_history, "skipped": False}

    # auto_rewrite のコア関数をインポート
    spec2 = importlib.util.spec_from_file_location("auto_rewrite", str(skill_scripts / "auto_rewrite.py"))
    auto_rw = importlib.util.module_from_spec(spec2)
    spec2.loader.exec_module(auto_rw)

    # ANTHROPIC_API_KEY チェック
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("  ANTHROPIC_API_KEY 未設定。AI臭除去のリライトをスキップ。")
        return {"ai_score": ai_score, "ai_rounds": 0, "ai_history": ai_history, "skipped": True}

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
    except ImportError:
        print("  anthropic パッケージ未インストール。AI臭除去のリライトをスキップ。")
        return {"ai_score": ai_score, "ai_rounds": 0, "ai_history": ai_history, "skipped": True}

    # 参照ファイル読み込み
    skill_root = skill_scripts.parent
    system_prompt = ""
    rewrite_prompt_path = skill_root / "prompts" / "rewrite_system.txt"
    if rewrite_prompt_path.exists():
        system_prompt = rewrite_prompt_path.read_text(encoding="utf-8")

    patterns_path = skill_root / "reference" / "ai_smell_patterns.md"
    patterns_text = patterns_path.read_text(encoding="utf-8") if patterns_path.exists() else ""

    good_examples_path = skill_root / "reference" / "good_examples.md"
    good_examples_text = good_examples_path.read_text(encoding="utf-8") if good_examples_path.exists() else ""

    vocab_path = skill_root / "reference" / "industry_vocab.json"
    import json
    vocab_data = json.loads(vocab_path.read_text(encoding="utf-8")) if vocab_path.exists() else {}

    full_system = f"{system_prompt}\n\n---\n\n## 参照: AI臭パターン辞典\n\n{patterns_text}\n\n---\n\n## 参照: 採択済み申請書の文体サンプル\n\n{good_examples_text}"

    # リライトループ
    current_text = text
    for round_num in range(1, max_rounds + 1):
        print(f"\n  AI臭除去 ラウンド {round_num}/{max_rounds}...")

        weak_areas = auto_rw.identify_weak_areas(result)
        instruction = auto_rw.build_rewrite_instruction(
            weak_areas, industry, round_num, vocab_data, None,
        )

        try:
            rewritten = auto_rw.rewrite_with_claude(
                client, current_text, full_system, instruction,
                auto_rw.DEFAULT_MODEL,
            )
        except Exception as e:
            print(f"  リライトAPI失敗: {e}")
            break

        result = ai_smell.calculate_score(rewritten)
        ai_score = result["total_score"]
        ai_history.append({"round": round_num, "score": ai_score, "grade": result["grade"]})
        print(f"  AI臭スコア（ラウンド{round_num}）: {ai_score}/100 ({result['grade']})")

        if on_progress:
            on_progress(f"ai_smell_round_{round_num}", ai_score, result)

        current_text = rewritten

        if ai_score >= target_ai_score:
            print(f"  AI臭スコア目標達成！ {ai_score} >= {target_ai_score}")
            break

        # スコアが改善しなかったら終了
        if round_num >= 2 and ai_history[-1]["score"] <= ai_history[-2]["score"]:
            print(f"  スコア改善なし。ループ終了。")
            break

    # リライト結果をdocxに書き戻し
    if len(ai_history) > 1:
        print(f"  リライト結果をdocxに書き戻し中...")
        _write_text_to_docx(output_dir, current_text)
        # リライト済みテキストも保存
        rewrite_path = Path(output_dir) / "事業計画書_リライト済み.txt"
        rewrite_path.write_text(current_text, encoding="utf-8")
        print(f"  保存: {rewrite_path}")

    return {"ai_score": ai_score, "ai_rounds": len(ai_history) - 1, "ai_history": ai_history, "skipped": False}


def generate_with_auto_fix(
    data: HearingData,
    output_dir: str,
    template_dir,
    diagrams: dict = None,
    target_score: int = 85,
    max_iterations: int = 5,
    skip_diagrams: bool = False,
    deai: bool = True,
    target_ai_score: int = 85,
    max_ai_rounds: int = 3,
    on_progress=None,
) -> dict:
    """スコアが目標に達するまで生成→検証→修正を繰り返し、
    品質スコア達成後にAI臭除去フェーズを実行する。

    Args:
        data: ヒアリングデータ
        output_dir: 出力ディレクトリ
        template_dir: テンプレートディレクトリ
        diagrams: 図解辞書（None=スキップ）
        target_score: 品質目標スコア（デフォルト85）
        max_iterations: 品質ループ最大リトライ回数（デフォルト5）
        skip_diagrams: 図解を採点から除外するか
        deai: AI臭除去フェーズを実行するか（デフォルトTrue）
        target_ai_score: AI臭スコア目標（デフォルト85）
        max_ai_rounds: AI臭除去の最大リライト回数（デフォルト3）
        on_progress: コールバック fn(phase, score, detail) — UIへの進捗通知用

    Returns:
        dict: {score, iterations, history, result, ai_result}
    """
    from validate import calculate_score

    if diagrams is None:
        diagrams = {}

    # 自動修正ループ開始前に成長率をデフォルト値にリセット
    # （前回のループで変更された値が残らないようにする）
    Config.reset_rates()

    history = []

    # === Phase 1: 書類品質ループ ===
    for iteration in range(1, max_iterations + 1):
        # --- 生成 ---
        _run_generation(data, output_dir, template_dir, diagrams)

        # --- テキスト穴あき修正（生成直後に実施）---
        hole_fixes = _fix_text_holes_in_docx(output_dir, data)
        if hole_fixes:
            print(f"  テキスト穴あき修正 {len(hole_fixes)}件:")
            for hf in hole_fixes:
                print(f"    - {hf}")

        # --- スコアリング ---
        result = calculate_score(Path(output_dir), skip_diagrams=skip_diagrams)
        current_score = result["score"]
        history.append({
            "iteration": iteration,
            "score": current_score,
            "breakdown": result["breakdown"],
            "issues": [i["detail"] for i in result["issues"]],
        })

        if on_progress:
            on_progress(iteration, current_score, history[-1])

        print(f"\n{'='*50}")
        print(f"  イテレーション {iteration}/{max_iterations}: 品質スコア {current_score}/100")
        for cat, info in result["breakdown"].items():
            print(f"    {cat}: {info['score']}/{info['max']}")

        # --- 目標達成チェック ---
        if current_score >= target_score:
            print(f"  品質スコア {target_score} を達成！")
            break

        # --- 最終イテレーションなら終了 ---
        if iteration >= max_iterations:
            print(f"  最大イテレーション {max_iterations} に到達。最終スコア: {current_score}")
            break

        # --- 自動修正 ---
        fixes = _apply_fixes(result["issues"], data)
        if not fixes:
            print(f"  追加の自動修正なし。最終スコア: {current_score}")
            break

        print(f"  自動修正を適用:")
        for fix in fixes:
            print(f"    - {fix}")

        # 出力ディレクトリをクリーンアップして再生成
        out_path = Path(output_dir)
        for f in out_path.glob("*_完成版.*"):
            f.unlink()

    # === Phase 2: AI臭除去 ===
    ai_result = {"ai_score": None, "ai_rounds": 0, "ai_history": [], "skipped": True}
    if deai:
        industry = data.company.industry or "サービス"
        print(f"\n{'='*50}")
        print(f"  Phase 2: AI臭除去（業種: {industry}）")
        ai_result = _run_deai_phase(
            output_dir=output_dir,
            industry=industry,
            target_ai_score=target_ai_score,
            max_rounds=max_ai_rounds,
            on_progress=on_progress,
        )

    # ループ終了後に成長率をデフォルトにリセット（他処理への影響防止）
    Config.reset_rates()

    final = calculate_score(Path(output_dir), skip_diagrams=skip_diagrams)
    return {
        "score": final["score"],
        "iterations": len(history),
        "history": history,
        "result": final,
        "ai_result": ai_result,
    }


# =============================================================================
# メイン処理
# =============================================================================

def main():
    import argparse

    parser = argparse.ArgumentParser(description="省力化補助金申請書類生成 v10.5 完全版")
    parser.add_argument("--hearing", "-H", required=False, help="ヒアリングシートのパス")
    parser.add_argument("--from-transcription", help="議事録テキストからヒアリングシートを自動生成して使用")
    parser.add_argument("--output", "-o", default="./output", help="出力ディレクトリ")
    parser.add_argument("--template-dir", "-t", required=True, help="テンプレートディレクトリ")
    parser.add_argument("--no-diagrams", action="store_true", help="図解生成をスキップ")
    parser.add_argument("--auto-fix", action="store_true", help="85点以上になるまで自動修正ループ")
    parser.add_argument("--target-score", type=int, default=85, help="自動修正の目標スコア（デフォルト85）")
    parser.add_argument("--max-iterations", type=int, default=5, help="自動修正の最大リトライ回数（デフォルト5）")
    parser.add_argument("--no-deai", action="store_true", help="AI臭除去フェーズをスキップ")
    parser.add_argument("--target-ai-score", type=int, default=85, help="AI臭除去の目標スコア（デフォルト85）")
    parser.add_argument("--max-ai-rounds", type=int, default=3, help="AI臭除去の最大リライト回数（デフォルト3）")
    args = parser.parse_args()

    # --hearing か --from-transcription のいずれかが必須
    if not args.hearing and not args.from_transcription:
        parser.error("--hearing または --from-transcription のいずれかを指定してください")

    print("=" * 70)
    print("省力化補助金 申請書類生成スクリプト v10.5 完全版")
    print("- 事業者概要テーブル完全対応")
    print("- PREP法による採択レベル文章生成")
    print("- nano-banana-pro-preview 図解生成")
    print("=" * 70)

    template_dir = Path(args.template_dir)
    output_dir = Path(args.output)
    output_dir.mkdir(exist_ok=True, parents=True)

    # 1. データ読み込み
    hearing_path = args.hearing
    if args.from_transcription:
        # 議事録テキストから一時Excelを生成
        import tempfile
        from transcription_to_hearing import transcription_to_hearing as t2h
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            print("❌ --from-transcription 使用時は ANTHROPIC_API_KEY 環境変数が必要です")
            sys.exit(1)
        tmp_hearing = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        tmp_hearing.close()
        _, _, hearing_path = t2h(
            input_path=args.from_transcription,
            output_path=tmp_hearing.name,
            api_key=api_key,
        )
        print(f"  📄 生成されたヒアリングシート: {hearing_path}")

    data = read_hearing_sheet(hearing_path)

    # 2. 図解生成
    diagrams = {} if args.no_diagrams else generate_diagrams(data, str(output_dir))

    if args.auto_fix:
        # 自動修正ループ
        deai_enabled = not args.no_deai
        print(f"\n🔄 自動修正モード: 品質目標 {args.target_score}点 / 最大 {args.max_iterations}回")
        if deai_enabled:
            print(f"   AI臭除去: 目標 {args.target_ai_score}点 / 最大 {args.max_ai_rounds}回")
        result = generate_with_auto_fix(
            data=data,
            output_dir=str(output_dir),
            template_dir=template_dir,
            diagrams=diagrams,
            target_score=args.target_score,
            max_iterations=args.max_iterations,
            skip_diagrams=args.no_diagrams,
            deai=deai_enabled,
            target_ai_score=args.target_ai_score,
            max_ai_rounds=args.max_ai_rounds,
        )
        print("\n" + "=" * 70)
        print(f"品質スコア: {result['score']}/100 （{result['iterations']}回で完了）")
        for h in result["history"]:
            status = "PASS" if h["score"] >= args.target_score else "----"
            print(f"  [{status}] #{h['iteration']}: {h['score']}点")
        ai_r = result.get("ai_result", {})
        if ai_r and not ai_r.get("skipped"):
            print(f"AI臭スコア: {ai_r['ai_score']}/100 （{ai_r['ai_rounds']}回リライト）")
            for ah in ai_r.get("ai_history", []):
                print(f"  ラウンド{ah['round']}: {ah['score']}点 ({ah['grade']})")
        print(f"📁 出力先: {output_dir}")
        print("=" * 70)
    else:
        # 通常の1回生成
        _run_generation(data, str(output_dir), template_dir, diagrams)
        print("\n" + "=" * 70)
        print("✅ 全ての書類生成が完了しました！")
        print(f"📁 出力先: {output_dir}")
        print("=" * 70)


if __name__ == "__main__":
    main()
