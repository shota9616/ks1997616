#!/usr/bin/env python3
"""自動修正ループ・AI臭除去"""

import os
import re
import json
from pathlib import Path

from docx import Document

from models import HearingData
from config import Config
from financial_utils import calc_base_components, calc_year_added_value, validate_financial_inputs, check_requirements, find_growth_rate_for_target_cagr
from document_writer import generate_business_plan_1_2
from plan3_writer import generate_business_plan_3
from other_documents import generate_other_documents


def _run_generation(data: HearingData, output_dir: str, template_dir, diagrams: dict):
    """書類一式を生成する（1回分の実行）"""
    template_dir = Path(template_dir)
    Path(output_dir).mkdir(exist_ok=True, parents=True)

    t = template_dir / "事業計画書_その1その2_様式.docx"
    if t.exists():
        generate_business_plan_1_2(data, diagrams, str(output_dir), t)

    t = template_dir / "事業計画書_その3_様式.xlsx"
    if t.exists():
        generate_business_plan_3(data, str(output_dir), t)

    generate_other_documents(data, str(output_dir), template_dir)


def _fix_text_holes_in_docx(output_dir: str, data: HearingData) -> list:
    """docx内のテキスト穴あき（プレースホルダー空白）を修正する"""
    docx_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return []

    doc = Document(str(docx_path))
    fixes = []

    # 修正マッピング: パターン → 置換テキスト生成関数
    e = data.equipment
    s = data.labor_shortage
    features = e.features if e.features else f"{e.name}による業務自動化・効率化機能"

    replacements = {
        "主要機能として、が挙げられる": f"主要機能として、{features}が挙げられる",
    }

    # 未丸め小数値の修正パターン
    decimal_pattern = re.compile(r"(\d+)\.(\d{6,})(\s*時間|\s*分)")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    text = original

                    # 固定パターン置換
                    for old_text, new_text in replacements.items():
                        if old_text in text:
                            text = text.replace(old_text, new_text)
                            fixes.append(f"穴あき修正: '{old_text}' → '{new_text[:40]}...'")

                    # 未丸め小数値修正
                    def _round_decimal(m):
                        whole = m.group(1)
                        decimal = m.group(2)
                        unit = m.group(3)
                        rounded = round(float(f"{whole}.{decimal}"), 1)
                        return f"{rounded}{unit}"
                    text = decimal_pattern.sub(_round_decimal, text)
                    if text != original:
                        if "小数丸め" not in str(fixes):
                            fixes.append("小数値丸め: 未整形の小数値を修正")

                    if text != original:
                        para.text = text

    for para in doc.paragraphs:
        original = para.text
        text = original
        for old_text, new_text in replacements.items():
            if old_text in text:
                text = text.replace(old_text, new_text)
                if f"穴あき修正: '{old_text}'" not in str(fixes):
                    fixes.append(f"穴あき修正: '{old_text}' → '{new_text[:40]}...'")
        text = decimal_pattern.sub(_round_decimal, text)
        if text != original:
            para.text = text

    if fixes:
        doc.save(str(docx_path))

    return fixes


def _fix_consistency_in_docx(output_dir: str, data: HearingData) -> list:
    """docx内の付加価値額をExcel参考書式と統一する（書類間整合性の自動修正）

    financial_utils.calc_base_components() を使用して「正しい付加価値額」を算出し、
    docx内の不一致値をrun単位で置換する。
    """
    docx_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return []

    base = calc_base_components(data)
    correct_av = base["added_value"]

    doc = Document(str(docx_path))
    fixes = []
    av_pattern = re.compile(r"(付加価値額[^0-9]{0,30}?(?:約)?)([0-9,]+)(円)")

    # テーブル内のセルを走査（事業計画書の本文はテーブル内に格納）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    m = av_pattern.search(full_text)
                    if m:
                        old_val = int(m.group(2).replace(",", ""))
                        if old_val != correct_av and old_val > 10000:
                            diff_ratio = abs(old_val - correct_av) / max(old_val, correct_av, 1)
                            if diff_ratio > 0.10:  # 10%以上の差異
                                old_str = f"{old_val:,}"
                                new_str = f"{correct_av:,}"
                                # run単位で置換（フォーマット保持）
                                for run in para.runs:
                                    if old_str in run.text:
                                        run.text = run.text.replace(old_str, new_str)
                                # runで置換できない場合（数字が複数runに分割）
                                if old_str in para.text:
                                    for run in para.runs:
                                        if old_str in run.text:
                                            run.text = run.text.replace(old_str, new_str)
                                fixes.append(f"付加価値額修正: {old_str}円 → {new_str}円")

    # 5年計画の数値も修正
    year_values = {yr: calc_year_added_value(base, yr) for yr in range(0, 6)}

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    # 「N年目：約XX,XXX,XXX円」パターン
                    year_pattern = re.compile(r"(\d)年目[：:]約?([0-9,]+)円")
                    for ym in year_pattern.finditer(text):
                        yr_num = int(ym.group(1))
                        old_yr_val = int(ym.group(2).replace(",", ""))
                        if yr_num in year_values:
                            new_yr_val = year_values[yr_num]
                            if old_yr_val != new_yr_val and abs(old_yr_val - new_yr_val) / max(old_yr_val, 1) > 0.10:
                                old_yr_str = f"{old_yr_val:,}"
                                new_yr_str = f"{new_yr_val:,}"
                                for run in para.runs:
                                    if old_yr_str in run.text:
                                        run.text = run.text.replace(old_yr_str, new_yr_str)
                                if old_yr_str not in [f.split("→")[0] for f in fixes]:
                                    fixes.append(f"計画値修正({yr_num}年目): {old_yr_str}円 → {new_yr_str}円")

                    # 「基準年度：約XX,XXX,XXX円」パターン
                    base_pattern = re.compile(r"基準年度[：:]約?([0-9,]+)円")
                    bm = base_pattern.search(text)
                    if bm:
                        old_base = int(bm.group(1).replace(",", ""))
                        if old_base != correct_av and abs(old_base - correct_av) / max(old_base, 1) > 0.10:
                            old_b_str = f"{old_base:,}"
                            new_b_str = f"{correct_av:,}"
                            for run in para.runs:
                                if old_b_str in run.text:
                                    run.text = run.text.replace(old_b_str, new_b_str)
                            fixes.append(f"基準年度修正: {old_b_str}円 → {new_b_str}円")

    if fixes:
        doc.save(str(docx_path))
        print(f"  📝 付加価値額整合性修正 {len(fixes)}件")

    return fixes


def _apply_fixes(issues: list, data: HearingData) -> list:
    """スコアリング結果のissuesを解析し、パラメータを自動修正する。
    適用した修正のリストを返す。"""
    fixes_applied = []

    for issue in issues:
        action = issue.get("action", "")

        if action == "increase_growth_rate":
            old = Config.GROWTH_RATE
            Config.GROWTH_RATE = min(Config.GROWTH_RATE + 0.005, 1.10)  # 上限10%
            if Config.GROWTH_RATE != old:
                fixes_applied.append(f"GROWTH_RATE: {old} -> {Config.GROWTH_RATE}")

        elif action == "increase_salary_rate":
            old = Config.SALARY_GROWTH_RATE
            Config.SALARY_GROWTH_RATE = min(Config.SALARY_GROWTH_RATE + 0.005, 1.08)  # 上限8%
            if Config.SALARY_GROWTH_RATE != old:
                fixes_applied.append(f"SALARY_GROWTH_RATE: {old} -> {Config.SALARY_GROWTH_RATE}")

        elif action == "increase_text" or action == "increase_section_text":
            # テキスト不足はテンプレートで対応済みのため、再生成で解決を試みる
            if "テキスト再生成" not in [f.split(":")[0] for f in fixes_applied]:
                fixes_applied.append("テキスト再生成: リトライ")

        elif action == "fix_text_holes":
            # docx直接編集で対応済み（Configパラメータ変更なし→再生成不要）
            print("    [INFO] テキスト穴あきはdocx直接編集で対応済み")

        elif action == "fix_value_inconsistency":
            # 整合性は同一Configでの再生成では改善しない→再生成不要
            print("    [INFO] 数値整合性は現在のConfig設定で統一済み")

        elif action == "fix_empty_section":
            # 空セクションはテンプレート構造の問題→再生成では改善しない
            print("    [INFO] 空セクションはテンプレート構造依存のため再生成不要")

        elif action == "fix_negative_profit":
            # 営業利益マイナスは成長率増加で対応
            old = Config.GROWTH_RATE
            Config.GROWTH_RATE = min(Config.GROWTH_RATE + 0.01, 1.10)
            if Config.GROWTH_RATE != old:
                fixes_applied.append(f"営業利益修正: GROWTH_RATE {old} -> {Config.GROWTH_RATE}")

    return fixes_applied


def _extract_docx_text(output_dir: str) -> str:
    """事業計画書docxから全テキストを抽出する"""
    docx_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return ""
    doc = Document(str(docx_path))
    texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        texts.append(t)
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            texts.append(t)
    return "\n\n".join(texts)


def _write_text_to_docx(output_dir: str, rewritten_text: str):
    """リライト済みテキストを事業計画書docxのテーブルセルに書き戻す"""
    docx_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return

    doc = Document(str(docx_path))

    # セクション番号→リライト済みテキストのマッピングを構築
    # リライト済みテキストをセクションヘッダーで分割
    section_map = {}
    current_key = None
    current_lines = []

    for line in rewritten_text.split("\n"):
        # セクションヘッダー検出（【...】パターン）
        header_match = re.match(r"^【(.+?)】", line.strip())
        if header_match:
            if current_key and current_lines:
                section_map[current_key] = "\n".join(current_lines).strip()
            current_key = header_match.group(1)
            current_lines = [line]
        elif current_key:
            current_lines.append(line)

    if current_key and current_lines:
        section_map[current_key] = "\n".join(current_lines).strip()

    if not section_map:
        # セクション分割できない場合、全体を最大のテーブルセルに書き込む
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if len(cell.text) > 500:
                        cell.text = rewritten_text
                        doc.save(str(docx_path))
                        return
        return

    # テーブルセルをスキャンし、対応するセクションのテキストを置換
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                for key, new_text in section_map.items():
                    if key in cell_text and len(cell_text) > 200:
                        cell.text = new_text
                        break

    doc.save(str(docx_path))


def _run_deai_phase(
    output_dir: str,
    industry: str,
    target_ai_score: int = 85,
    max_rounds: int = 3,
    on_progress=None,
) -> dict:
    """AI臭除去フェーズ: docxテキスト抽出→スコアリング→リライト→書き戻し

    Returns:
        dict: {ai_score, ai_rounds, ai_history, skipped}
    """
    # ai_smell_score をインポート
    skill_scripts = Path.home() / ".claude" / "skills" / "shoryokuka-review-deai" / "scripts"
    if not skill_scripts.exists():
        print("  AI臭除去スキルが未インストール。スキップします。")
        return {"ai_score": None, "ai_rounds": 0, "ai_history": [], "skipped": True}

    import importlib.util
    spec = importlib.util.spec_from_file_location("ai_smell_score", str(skill_scripts / "ai_smell_score.py"))
    ai_smell = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(ai_smell)

    # テキスト抽出
    text = _extract_docx_text(output_dir)
    if not text or len(text) < 100:
        print("  事業計画書テキストが短すぎます。AI臭除去をスキップ。")
        return {"ai_score": None, "ai_rounds": 0, "ai_history": [], "skipped": True}

    # 初回スコアリング
    result = ai_smell.calculate_score(text)
    ai_score = result["total_score"]
    ai_history = [{"round": 0, "score": ai_score, "grade": result["grade"]}]
    print(f"\n  AI臭スコア（初回）: {ai_score}/100 ({result['grade']})")

    if on_progress:
        on_progress("ai_smell_initial", ai_score, result)

    if ai_score >= target_ai_score:
        print(f"  AI臭スコア {ai_score} >= {target_ai_score}。リライト不要。")
        return {"ai_score": ai_score, "ai_rounds": 0, "ai_history": ai_history, "skipped": False}

    # auto_rewrite のコア関数をインポート
    spec2 = importlib.util.spec_from_file_location("auto_rewrite", str(skill_scripts / "auto_rewrite.py"))
    auto_rw = importlib.util.module_from_spec(spec2)
    spec2.loader.exec_module(auto_rw)

    # ANTHROPIC_API_KEY チェック
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("  ANTHROPIC_API_KEY 未設定。AI臭除去のリライトをスキップ。")
        return {"ai_score": ai_score, "ai_rounds": 0, "ai_history": ai_history, "skipped": True}

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
    except ImportError:
        print("  anthropic パッケージ未インストール。AI臭除去のリライトをスキップ。")
        return {"ai_score": ai_score, "ai_rounds": 0, "ai_history": ai_history, "skipped": True}

    # 参照ファイル読み込み
    skill_root = skill_scripts.parent
    system_prompt = ""
    rewrite_prompt_path = skill_root / "prompts" / "rewrite_system.txt"
    if rewrite_prompt_path.exists():
        system_prompt = rewrite_prompt_path.read_text(encoding="utf-8")

    patterns_path = skill_root / "reference" / "ai_smell_patterns.md"
    patterns_text = patterns_path.read_text(encoding="utf-8") if patterns_path.exists() else ""

    good_examples_path = skill_root / "reference" / "good_examples.md"
    good_examples_text = good_examples_path.read_text(encoding="utf-8") if good_examples_path.exists() else ""

    vocab_path = skill_root / "reference" / "industry_vocab.json"
    vocab_data = json.loads(vocab_path.read_text(encoding="utf-8")) if vocab_path.exists() else {}

    full_system = f"{system_prompt}\n\n---\n\n## 参照: AI臭パターン辞典\n\n{patterns_text}\n\n---\n\n## 参照: 採択済み申請書の文体サンプル\n\n{good_examples_text}"

    # リライトループ
    current_text = text
    for round_num in range(1, max_rounds + 1):
        print(f"\n  AI臭除去 ラウンド {round_num}/{max_rounds}...")

        weak_areas = auto_rw.identify_weak_areas(result)
        instruction = auto_rw.build_rewrite_instruction(
            weak_areas, industry, round_num, vocab_data, None,
        )

        try:
            rewritten = auto_rw.rewrite_with_claude(
                client, current_text, full_system, instruction,
                auto_rw.DEFAULT_MODEL,
            )
        except Exception as e:
            print(f"  リライトAPI失敗: {e}")
            break

        result = ai_smell.calculate_score(rewritten)
        ai_score = result["total_score"]
        ai_history.append({"round": round_num, "score": ai_score, "grade": result["grade"]})
        print(f"  AI臭スコア（ラウンド{round_num}）: {ai_score}/100 ({result['grade']})")

        if on_progress:
            on_progress(f"ai_smell_round_{round_num}", ai_score, result)

        current_text = rewritten

        if ai_score >= target_ai_score:
            print(f"  AI臭スコア目標達成！ {ai_score} >= {target_ai_score}")
            break

        # スコアが改善しなかったら終了
        if round_num >= 2 and ai_history[-1]["score"] <= ai_history[-2]["score"]:
            print(f"  スコア改善なし。ループ終了。")
            break

    # リライト結果をdocxに書き戻し
    if len(ai_history) > 1:
        print(f"  リライト結果をdocxに書き戻し中...")
        _write_text_to_docx(output_dir, current_text)
        # リライト済みテキストも保存
        rewrite_path = Path(output_dir) / "事業計画書_リライト済み.txt"
        rewrite_path.write_text(current_text, encoding="utf-8")
        print(f"  保存: {rewrite_path}")

    return {"ai_score": ai_score, "ai_rounds": len(ai_history) - 1, "ai_history": ai_history, "skipped": False}


def generate_with_auto_fix(
    data: HearingData,
    output_dir: str,
    template_dir,
    diagrams: dict = None,
    target_score: int = 85,
    max_iterations: int = 5,
    skip_diagrams: bool = False,
    deai: bool = True,
    target_ai_score: int = 85,
    max_ai_rounds: int = 3,
    on_progress=None,
) -> dict:
    """スコアが目標に達するまで生成→検証→修正を繰り返し、
    品質スコア達成後にAI臭除去フェーズを実行する。
    """
    from validate import calculate_score

    if diagrams is None:
        diagrams = {}

    # 自動修正ループ開始前に成長率をデフォルト値にリセット
    # （前回のループで変更された値が残らないようにする）
    Config.reset_rates()

    history = []

    # === Phase 0: 事前バリデーションゲート（★追加）===
    input_warnings = validate_financial_inputs(data)
    if input_warnings:
        print("\n⚠️ 入力データ事前チェック警告:")
        for w in input_warnings:
            print(f"  {w}")
        if on_progress:
            on_progress("pre_validation", 0, {"warnings": input_warnings})

    # 要件チェック（生成前に成長率パラメータが要件を満たすか確認）
    req_check = check_requirements(data)
    if req_check["warnings"]:
        print("\n⚠️ 要件充足チェック警告（生成前）:")
        for w in req_check["warnings"]:
            print(f"  {w}")
        # SALARY_GROWTH_RATEが1人当たり要件を満たさない場合は自動調整
        if not req_check["salary_per_capita_ok"]:
            # 要件を満たす最小成長率を逆算: (1+r)^5 >= (1+0.035)^5
            min_rate = 1 + Config.REQUIREMENT_SALARY_PER_CAPITA_CAGR
            if Config.SALARY_GROWTH_RATE < min_rate:
                old_rate = Config.SALARY_GROWTH_RATE
                Config.SALARY_GROWTH_RATE = min_rate + 0.005  # 余裕を持たせる
                print(f"  🔧 SALARY_GROWTH_RATE自動調整: {old_rate} → {Config.SALARY_GROWTH_RATE}")

        # 労働生産性CAGR（≧4.0%）が未達の場合、GROWTH_RATEを逆算して引き上げ
        if not req_check.get("labor_productivity_ok", True):
            # 減価償却費（成長率0%）の構成比に応じて必要なGROWTH_RATEが異なるため、
            # 企業データから逆算する
            required_rate = find_growth_rate_for_target_cagr(
                data, Config.REQUIREMENT_LABOR_PRODUCTIVITY_CAGR
            )
            if Config.GROWTH_RATE < required_rate:
                old_rate = Config.GROWTH_RATE
                Config.GROWTH_RATE = required_rate
                print(f"  🔧 GROWTH_RATE自動調整（労働生産性要件）: {old_rate:.3f} → {Config.GROWTH_RATE:.3f}")

    # === Phase 1: 書類品質ループ ===
    for iteration in range(1, max_iterations + 1):
        # --- 生成 ---
        _run_generation(data, output_dir, template_dir, diagrams)

        # --- テキスト穴あき修正（生成直後に実施）---
        hole_fixes = _fix_text_holes_in_docx(output_dir, data)
        if hole_fixes:
            print(f"  テキスト穴あき修正 {len(hole_fixes)}件:")
            for hf in hole_fixes:
                print(f"    - {hf}")

        # --- 付加価値額の書類間整合性修正 ---
        consistency_fixes = _fix_consistency_in_docx(output_dir, data)
        if consistency_fixes:
            for cf in consistency_fixes:
                print(f"    - {cf}")

        # --- スコアリング ---
        result = calculate_score(Path(output_dir), skip_diagrams=skip_diagrams)
        current_score = result["score"]
        history.append({
            "iteration": iteration,
            "score": current_score,
            "breakdown": result["breakdown"],
            "issues": [i["detail"] for i in result["issues"]],
        })

        if on_progress:
            on_progress(iteration, current_score, history[-1])

        print(f"\n{'='*50}")
        print(f"  イテレーション {iteration}/{max_iterations}: 品質スコア {current_score}/100")
        for cat, info in result["breakdown"].items():
            print(f"    {cat}: {info['score']}/{info['max']}")

        # --- 目標達成チェック ---
        if current_score >= target_score:
            print(f"  品質スコア {target_score} を達成！")
            break

        # --- スコア停滞検出（±1点以内を停滞と判定）---
        if iteration > 1 and abs(current_score - history[-2]["score"]) <= 1.0:
            print(f"  スコア停滞を検出（{current_score}点 ≒ {history[-2]['score']}点）。再生成しても改善しないためループ終了。")
            break

        # --- 最終イテレーションなら終了 ---
        if iteration >= max_iterations:
            print(f"  最大イテレーション {max_iterations} に到達。最終スコア: {current_score}")
            break

        # --- 自動修正 ---
        fixes = _apply_fixes(result["issues"], data)
        if not fixes:
            print(f"  追加の自動修正なし。最終スコア: {current_score}")
            break

        print(f"  自動修正を適用:")
        for fix in fixes:
            print(f"    - {fix}")

        # 出力ディレクトリをクリーンアップして再生成
        out_path = Path(output_dir)
        for f in out_path.glob("*_完成版.*"):
            f.unlink()

    # === Phase 2: AI臭除去 ===
    ai_result = {"ai_score": None, "ai_rounds": 0, "ai_history": [], "skipped": True}
    if deai:
        industry = data.company.industry or "サービス"
        print(f"\n{'='*50}")
        print(f"  Phase 2: AI臭除去（業種: {industry}）")
        ai_result = _run_deai_phase(
            output_dir=output_dir,
            industry=industry,
            target_ai_score=target_ai_score,
            max_rounds=max_ai_rounds,
            on_progress=on_progress,
        )

    # ループ終了後に成長率をデフォルトにリセット（他処理への影響防止）
    Config.reset_rates()

    final = calculate_score(Path(output_dir), skip_diagrams=skip_diagrams)
    return {
        "score": final["score"],
        "iterations": len(history),
        "history": history,
        "result": final,
        "ai_result": ai_result,
    }
