#!/usr/bin/env python3
"""
省力化補助金 申請書類 出力検証スクリプト

生成された書類の品質を自動チェックする。
- ファイル存在確認（11種）
- 事業計画書の文字数チェック（総合＋セクション別）
- 図解ファイル数チェック（11種）
- 基本要件の数値整合性チェック
- JSON形式でのレポート出力（オプション）
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from docx import Document
except ImportError:
    Document = None


# =============================================================================
# 定数
# =============================================================================

EXPECTED_FILES = [
    "事業計画書_その1その2_完成版.docx",
    "事業計画書_その3_完成版.xlsx",
    "役員名簿_完成版.xlsx",
    "従業員名簿_完成版.xlsx",
    "株主出資者名簿_完成版.xlsx",
    "事業実施場所リスト_完成版.xlsx",
    "他の補助金使用実績_完成版.xlsx",
    "金融機関確認書_完成版.docx",
    "給与支給総額確認書_完成版.xlsx",
    "賃金引上げ要件_事業場内_完成版.xlsx",
    "賃金引上げ要件_地域別_完成版.xlsx",
]

SECTION_HEADERS = {
    "1-1": ["現状分析", "事業の現状"],
    "1-2": ["経営課題", "人手不足", "経営上の課題"],
    "1-3": ["動機目的", "動機", "なぜ今", "省力化補助金活用の動機"],
    "2-1": ["ビフォーアフター", "導入前後", "省力化の内容", "省力化設備導入による業務プロセス"],
    "2-2": ["効果", "省力化効果", "省力化投資により期待される効果"],
    "3-1": ["生産性向上", "賃上げ", "事業場内", "付加価値額の向上"],
}

MIN_CHAR_COUNTS = {
    "1-1 現状分析": 600,
    "1-2 経営課題": 700,
    "1-3 動機目的": 400,
    "2-1 ビフォーアフター": 1000,
    "2-2 効果": 600,
    "3-1 生産性向上": 700,
}

MIN_TOTAL_CHARS = 4700

EXPECTED_DIAGRAM_COUNT = 13

# テキスト品質チェック: 穴あき・テンプレ残留パターン
TEXT_HOLE_PATTERNS = [
    (r"、が挙げられる", "プレースホルダー空白（features未入力）"),
    (r"、が期待できる", "プレースホルダー空白"),
    (r"として、を", "プレースホルダー空白"),
    (r"具体的には、に", "プレースホルダー空白"),
    (r"として、が", "プレースホルダー空白"),
    (r"主要機能として、が", "設備機能が未入力"),
    (r"\d+\.\d{6,}", "未丸め小数値（例: 4.083333...）"),
    (r"〇〇|●●|△△|□□|※※", "プレースホルダー記号の残留"),
    (r"None円|None名|None時間", "None値の残留"),
    (r"(?<![1-9,])0円[^0-9]*(?<![1-9,])0円[^0-9]*(?<![1-9,])0円", "全て0円（財務データ未入力）"),
]


# =============================================================================
# チェック関数
# =============================================================================

def check_files(output_dir: Path) -> list:
    """全11ファイルの存在確認"""
    results = []
    for fname in EXPECTED_FILES:
        fpath = output_dir / fname
        exists = fpath.exists()
        size = fpath.stat().st_size if exists else 0
        results.append({
            "file": fname,
            "exists": exists,
            "size": size,
            "ok": exists and size > 0,
        })
    return results


def check_diagrams(output_dir: Path) -> dict:
    """図解PNGファイルの数チェック"""
    diagram_dir = output_dir / "diagrams"
    if not diagram_dir.exists():
        return {
            "found": 0,
            "expected": EXPECTED_DIAGRAM_COUNT,
            "files": [],
            "ok": False,
            "note": "diagrams/ ディレクトリが存在しない",
        }

    png_files = sorted(diagram_dir.glob("*.png"))
    count = len(png_files)
    return {
        "found": count,
        "expected": EXPECTED_DIAGRAM_COUNT,
        "files": [f.name for f in png_files],
        "ok": count >= EXPECTED_DIAGRAM_COUNT,
    }


def _strip_whitespace(text: str) -> str:
    """空白・改行・タブを除去して文字数カウント用のテキストを返す"""
    return text.replace(" ", "").replace("\u3000", "").replace("\n", "").replace("\t", "")


def _normalize_section_number(text: str) -> str:
    """全角数字・ダッシュを半角に正規化"""
    table = str.maketrans("０１２３４５６７８９−‐‑‒–—―", "0123456789-------")
    return text.translate(table)


def _identify_section(text: str) -> str:
    """段落テキストからセクションIDを判定（ヘッダー検出用）
    テキストの冒頭にセクション番号が出現するかで判定する。
    テンプレート注釈内の参照（※1-2経営課題 等）は除外。"""
    normalized = _normalize_section_number(text)
    # 冒頭30文字以内にセクション番号があるかチェック（注釈内参照を除外）
    head = normalized[:30]
    for section_id, keywords in SECTION_HEADERS.items():
        # セクション番号が冒頭付近にある
        pattern = rf"^[^\w]*{re.escape(section_id)}"
        if re.search(pattern, head):
            return section_id
    # キーワードのみマッチ（短いヘッダー行）
    if len(_strip_whitespace(text)) < 60:
        for section_id, keywords in SECTION_HEADERS.items():
            for kw in keywords:
                if kw in text:
                    return section_id
    return ""


def check_docx_text(output_dir: Path) -> dict:
    """事業計画書docxの文字数チェック（総合＋セクション別）"""
    if Document is None:
        return {"error": "python-docxが未インストール"}

    docx_path = output_dir / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return {"error": "ファイルが存在しない"}

    doc = Document(str(docx_path))

    # 全テキスト収集
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text

    total_chars = len(_strip_whitespace(full_text))

    # セクション別文字数カウント（Phase 7: テーブルセル内テキストも走査）
    section_chars = {}
    current_section = ""
    section_texts = {}

    # パラグラフとテーブルセル内テキストの両方を走査
    all_texts = []
    for para in doc.paragraphs:
        all_texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_texts.append(para.text.strip())

    for text in all_texts:
        if not text:
            continue

        detected = _identify_section(text)
        if detected:
            current_section = detected
            if current_section not in section_texts:
                section_texts[current_section] = ""
            # ヘッダー行にコンテンツが含まれる場合（テンプレート注釈+本文が同一セル）
            # ヘッダーキーワード以降のテキストもカウントする
            if len(_strip_whitespace(text)) > 80:
                section_texts[current_section] += text
            continue

        if current_section and current_section in SECTION_HEADERS:
            section_texts.setdefault(current_section, "")
            section_texts[current_section] += text

    for section_id, text in section_texts.items():
        char_count = len(_strip_whitespace(text))
        # MIN_CHAR_COUNTSのキーとマッチさせる
        for key, min_count in MIN_CHAR_COUNTS.items():
            if key.startswith(section_id):
                section_chars[key] = {
                    "chars": char_count,
                    "min_required": min_count,
                    "ok": char_count >= min_count,
                }
                break

    return {
        "total_chars": total_chars,
        "min_required": MIN_TOTAL_CHARS,
        "ok": total_chars >= MIN_TOTAL_CHARS,
        "sections": section_chars,
    }


def check_text_quality(output_dir: Path) -> dict:
    """テキスト品質チェック: 穴あき・テンプレ残留・空セクション検出"""
    if Document is None:
        return {"error": "python-docxが未インストール", "issues": [], "score_ratio": 1.0}

    docx_path = output_dir / "事業計画書_その1その2_完成版.docx"
    if not docx_path.exists():
        return {"error": "ファイルが存在しない", "issues": [], "score_ratio": 0.0}

    doc = Document(str(docx_path))

    # 全テキスト収集
    all_texts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            all_texts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        all_texts.append(t)

    full_text = "\n".join(all_texts)
    issues = []

    # パターンマッチで穴あき検出
    for pattern, description in TEXT_HOLE_PATTERNS:
        matches = re.findall(pattern, full_text)
        if matches:
            issues.append({
                "type": "text_hole",
                "pattern": pattern,
                "description": description,
                "count": len(matches),
                "examples": matches[:3],
            })

    # 空セクション検出（見出しの直後にテキストがない）
    # docxではセクション見出し+テンプレート注釈+本文が同一セルに入ることがある
    for i, text in enumerate(all_texts):
        normalized_head = _normalize_section_number(text[:30])
        # セクション見出し行を検出
        if re.match(r"^[0-9][-.][0-9]", normalized_head):
            # ヘッダー行自体にコンテンツが含まれている場合（テンプレート注釈+本文が同一セル）
            # → テンプレート注釈（※）を除去した後の実質テキストが十分あれば空ではない
            text_without_annotation = re.sub(r"※[^\n]*", "", text)
            # セクション番号行自体を除去
            lines = text_without_annotation.split("\n")
            content_lines = [l for l in lines[1:] if _strip_whitespace(l)]
            content_in_header = sum(len(_strip_whitespace(l)) for l in content_lines)

            if content_in_header > 100:
                # 同一セル内に十分なコンテンツあり → 空セクションではない
                continue

            # 後続テキストにコンテンツがあるか確認
            has_content = False
            for j in range(1, min(6, len(all_texts) - i)):
                candidate = all_texts[i + j]
                candidate_norm = _normalize_section_number(candidate[:30])
                # 次のセクション見出しに到達 → コンテンツなし
                if re.match(r"^[0-9][-.][0-9]", candidate_norm):
                    break
                candidate_stripped = _strip_whitespace(candidate)
                # テンプレート注釈や短い行はスキップ
                if candidate.lstrip().startswith("※") or len(candidate_stripped) < 30:
                    continue
                has_content = True
                break
            if not has_content:
                issues.append({
                    "type": "empty_section",
                    "description": f"空セクション: {text[:30]}",
                    "header": text[:50],
                })

    # スコア比率計算（問題がないほど1.0に近い）
    penalty = len([i for i in issues if i["type"] == "text_hole"]) * 0.1
    penalty += len([i for i in issues if i["type"] == "empty_section"]) * 0.2
    score_ratio = max(0.0, 1.0 - penalty)

    return {
        "issues": issues,
        "issue_count": len(issues),
        "score_ratio": round(score_ratio, 2),
        "ok": len(issues) == 0,
    }


def check_cross_document_consistency(output_dir: Path) -> dict:
    """3書類間（docx/指定様式/参考書式）の数値整合性チェック"""
    if openpyxl is None or Document is None:
        return {"error": "必要パッケージ未インストール", "issues": [], "score_ratio": 1.0}

    issues = []

    # docxから付加価値額を抽出
    docx_path = output_dir / "事業計画書_その1その2_完成版.docx"
    docx_added_value = None
    if docx_path.exists():
        doc = Document(str(docx_path))
        for para in doc.paragraphs:
            m = re.search(r"付加価値額[^0-9]*([0-9,]+)円", para.text.replace("、", ""))
            if m:
                docx_added_value = int(m.group(1).replace(",", ""))
                break
        if docx_added_value is None:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        m = re.search(r"付加価値額[^0-9]*([0-9,]+)円", cell.text.replace("、", ""))
                        if m:
                            docx_added_value = int(m.group(1).replace(",", ""))
                            break

    # 参考書式シートから付加価値額を抽出
    xlsx_path = output_dir / "事業計画書_その3_完成版.xlsx"
    ref_added_value = None
    designated_added_value = None

    if xlsx_path.exists():
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            for name in wb.sheetnames:
                if "参考書式" in name or "目標値" in name:
                    ws = wb[name]
                    # 「付加価値額」ラベルを検索して対応するE列の値を取得
                    # （テンプレートによって行番号が変わるため動的検索）
                    for row_num in range(20, 50):
                        label_b = str(ws[f'B{row_num}'].value or "")
                        label_c = str(ws[f'C{row_num}'].value or "")
                        if "付加価値額" in label_b or "付加価値額" in label_c:
                            ref_added_value = ws[f'E{row_num}'].value
                            break
                    break
            for name in wb.sheetnames:
                if "指定様式" in name:
                    ws = wb[name]
                    # 付加価値額の基準年度値を探索
                    # 「①付加価値額」のように単独で付加価値額を指すセルを優先
                    for row in ws.iter_rows(min_row=1, max_row=60):
                        for cell in row:
                            cell_str = str(cell.value) if cell.value else ""
                            # 「付加価値額」を含み、かつ短いラベルセル（一覧行は除外）
                            if "付加価値額" in cell_str and len(cell_str) < 30:
                                # 同じ行の数値を取得（小さすぎる値は省力化指数等なので除外）
                                for c2 in row:
                                    if isinstance(c2.value, (int, float)) and c2.value > 10000:
                                        designated_added_value = c2.value
                                        break
                                if designated_added_value is not None:
                                    break
                        if designated_added_value is not None:
                            break
                    break

            # 参考書式の営業利益チェック（マイナスは致命的）
            for name in wb.sheetnames:
                if "参考書式" in name or "目標値" in name:
                    ws = wb[name]
                    # 営業利益セル（一般的な位置）
                    for col_letter in ['E', 'F', 'G', 'H', 'I', 'J', 'K']:
                        for row_num in [20, 21, 22, 23, 24, 25]:
                            val = ws[f'{col_letter}{row_num}'].value
                            if isinstance(val, (int, float)) and val < -1000000:
                                issues.append({
                                    "type": "negative_profit",
                                    "description": f"参考書式の営業利益がマイナス（{col_letter}{row_num}: {val:,.0f}円）",
                                    "cell": f"{col_letter}{row_num}",
                                    "value": val,
                                })
                    break

            wb.close()
        except Exception:
            pass

    # 3書類間の付加価値額の不整合チェック
    values = {}
    if docx_added_value is not None:
        values["docx本文"] = docx_added_value
    if ref_added_value is not None and isinstance(ref_added_value, (int, float)):
        values["参考書式"] = int(ref_added_value)
    if designated_added_value is not None:
        values["指定様式"] = int(designated_added_value)

    if len(values) >= 2:
        vals = list(values.values())
        max_val = max(vals)
        min_val = min(vals)
        if max_val > 0 and (max_val - min_val) / max_val > 0.15:  # 15%以上の差異
            issues.append({
                "type": "value_inconsistency",
                "description": f"付加価値額が書類間で不整合: {values}",
                "values": values,
            })

    penalty = 0
    for issue in issues:
        if issue["type"] == "value_inconsistency":
            penalty += 0.4
        elif issue["type"] == "negative_profit":
            penalty += 0.3
    score_ratio = max(0.0, 1.0 - penalty)

    return {
        "issues": issues,
        "issue_count": len(issues),
        "score_ratio": round(score_ratio, 2),
        "ok": len(issues) == 0,
    }


def check_plan3_values(output_dir: Path) -> dict:
    """事業計画書その3の数値チェック"""
    if openpyxl is None:
        return {"error": "openpyxlが未インストール"}

    xlsx_path = output_dir / "事業計画書_その3_完成版.xlsx"
    if not xlsx_path.exists():
        return {"error": "ファイルが存在しない"}

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception as e:
        return {"error": f"読み込みエラー: {e}"}

    results = {}

    # 別紙1チェック
    for name in wb.sheetnames:
        if "別紙1" in name or "省力化" in name:
            ws = wb[name]
            before_count = sum(1 for r in range(11, 17) if ws[f'C{r}'].value)
            after_count = sum(1 for r in range(11, 17) if ws[f'I{r}'].value)
            results["別紙1"] = {
                "導入前工程数": before_count,
                "導入後工程数": after_count,
                "ok": before_count > 0 and after_count > 0,
            }
            break

    # 参考書式チェック
    for name in wb.sheetnames:
        if "参考書式" in name or "目標値" in name:
            ws = wb[name]
            base_revenue = ws['E26'].value
            year5_revenue = ws['K26'].value
            base_salary = ws['E44'].value
            year5_salary = ws['K44'].value

            if base_revenue and year5_revenue and base_revenue > 0:
                revenue_growth = ((year5_revenue / base_revenue) ** (1/5) - 1) * 100
                results["付加価値額成長率"] = {
                    "年率": f"{revenue_growth:.1f}%",
                    "基準": "4.0%以上",
                    "ok": round(revenue_growth, 1) >= 4.0,
                }

            if base_salary and year5_salary and base_salary > 0:
                salary_growth = ((year5_salary / base_salary) ** (1/5) - 1) * 100
                results["給与支給総額成長率"] = {
                    "年率": f"{salary_growth:.1f}%",
                    "基準": "2.0%以上",
                    "ok": round(salary_growth, 1) >= 2.0,
                }
            break

    wb.close()
    return results


# =============================================================================
# スコアリング
# =============================================================================

def calculate_score(output_dir: Path, skip_diagrams: bool = False) -> dict:
    """出力書類を100点満点でスコアリングし、内訳と改善ヒントを返す。

    配点（審査基準に準拠した重み付け）:
      - ファイル存在:     10点 （形式要件：全書類生成の前提確認）
      - 図解:             5点  （補足資料）
      - テキスト量:       15点 （総5 + セクション別10）
      - テキスト品質:     25点 （穴あき・テンプレ残留・空セクション検出）★NEW
      - 数値整合性:       25点 （成長率15 + 書類間整合10）★強化
      - 数値要件:         20点 （付加価値額成長率 + 給与支給総額成長率）
    """
    output_dir = Path(output_dir)
    file_results = check_files(output_dir)
    diagram_results = check_diagrams(output_dir)
    text_results = check_docx_text(output_dir)
    value_results = check_plan3_values(output_dir)
    quality_results = check_text_quality(output_dir)
    consistency_results = check_cross_document_consistency(output_dir)

    score = 0.0
    breakdown = {}
    issues = []

    # --- ファイル存在 (10点) ---
    file_ok = sum(1 for r in file_results if r["ok"])
    file_total = len(file_results)
    file_score = (file_ok / file_total) * 10 if file_total > 0 else 0
    breakdown["files"] = {"score": round(file_score, 1), "max": 10, "detail": f"{file_ok}/{file_total}"}
    score += file_score
    if file_ok < file_total:
        missing = [r["file"] for r in file_results if not r["ok"]]
        issues.append({"category": "files", "action": "retry_generation", "detail": f"未生成ファイル: {', '.join(missing)}"})

    # --- 図解 (5点、skip_diagrams時は満点扱い) ---
    if skip_diagrams:
        diagram_score = 5.0
    else:
        d_found = diagram_results.get("found", 0)
        d_expected = diagram_results.get("expected", EXPECTED_DIAGRAM_COUNT)
        diagram_score = min(d_found / d_expected, 1.0) * 5 if d_expected > 0 else 0
        if d_found < d_expected:
            issues.append({"category": "diagrams", "action": "retry_diagrams", "detail": f"図解不足: {d_found}/{d_expected}"})
    breakdown["diagrams"] = {"score": round(diagram_score, 1), "max": 5, "detail": f"{diagram_results.get('found', 0)}/{diagram_results.get('expected', EXPECTED_DIAGRAM_COUNT)}"}
    score += diagram_score

    # --- 総文字数 (5点) ---
    if "error" not in text_results:
        total_chars = text_results.get("total_chars", 0)
        text_ratio = min(total_chars / MIN_TOTAL_CHARS, 1.0)
        text_score = text_ratio * 5
        if total_chars < MIN_TOTAL_CHARS:
            issues.append({"category": "text_total", "action": "increase_text", "detail": f"総文字数不足: {total_chars}/{MIN_TOTAL_CHARS}"})
    else:
        text_score = 0
        issues.append({"category": "text_total", "action": "retry_generation", "detail": text_results.get("error", "")})
    breakdown["text_total"] = {"score": round(text_score, 1), "max": 5, "detail": f"{text_results.get('total_chars', 0)}字"}
    score += text_score

    # --- セクション別文字数 (10点: 各~1.67点×6) ---
    section_score = 0
    per_section = 10.0 / len(MIN_CHAR_COUNTS)
    sections = text_results.get("sections", {}) if "error" not in text_results else {}
    for key, min_count in MIN_CHAR_COUNTS.items():
        sec = sections.get(key, {})
        chars = sec.get("chars", 0)
        ratio = min(chars / min_count, 1.0) if min_count > 0 else 1.0
        section_score += ratio * per_section
        if chars < min_count:
            issues.append({"category": "section", "action": "increase_section_text", "detail": f"{key}: {chars}/{min_count}字", "section": key})
    breakdown["sections"] = {"score": round(section_score, 1), "max": 10, "detail": f"{len([s for s in sections.values() if s.get('ok')])}/{len(MIN_CHAR_COUNTS)} セクション合格"}
    score += section_score

    # --- テキスト品質 (25点) ★NEW ---
    if "error" not in quality_results:
        quality_ratio = quality_results.get("score_ratio", 1.0)
        quality_score = quality_ratio * 25
        for qi in quality_results.get("issues", []):
            if qi["type"] == "text_hole":
                issues.append({
                    "category": "text_quality",
                    "action": "fix_text_holes",
                    "detail": f"テキスト穴あき: {qi['description']} ({qi['count']}箇所)",
                    "pattern": qi.get("pattern", ""),
                })
            elif qi["type"] == "empty_section":
                issues.append({
                    "category": "text_quality",
                    "action": "fix_empty_section",
                    "detail": f"空セクション: {qi.get('header', '')}",
                })
    else:
        quality_score = 0
    breakdown["text_quality"] = {"score": round(quality_score, 1), "max": 25, "detail": f"品質比率: {quality_results.get('score_ratio', 0):.0%}"}
    score += quality_score

    # --- 数値要件 (20点: 付加価値額10点 + 給与10点) ---
    value_score = 0
    if isinstance(value_results, dict) and "error" not in value_results:
        if "付加価値額成長率" in value_results:
            v = value_results["付加価値額成長率"]
            if v.get("ok"):
                value_score += 10
            else:
                rate_str = v.get("年率", "0%").replace("%", "")
                try:
                    rate = float(rate_str)
                    value_score += min(rate / 4.0, 1.0) * 10
                except ValueError:
                    pass
                issues.append({"category": "growth_rate", "action": "increase_growth_rate", "detail": f"付加価値額成長率: {v.get('年率', '?')}"})
        if "給与支給総額成長率" in value_results:
            v = value_results["給与支給総額成長率"]
            if v.get("ok"):
                value_score += 10
            else:
                rate_str = v.get("年率", "0%").replace("%", "")
                try:
                    rate = float(rate_str)
                    value_score += min(rate / 2.0, 1.0) * 10
                except ValueError:
                    pass
                issues.append({"category": "salary_rate", "action": "increase_salary_rate", "detail": f"給与支給総額成長率: {v.get('年率', '?')}"})
    else:
        issues.append({"category": "values", "action": "retry_generation", "detail": "数値チェック不能"})
    breakdown["values"] = {"score": round(value_score, 1), "max": 20}
    score += value_score

    # --- 書類間数値整合性 (25点) ★NEW ---
    if "error" not in consistency_results:
        consistency_ratio = consistency_results.get("score_ratio", 1.0)
        consistency_score = consistency_ratio * 25
        for ci in consistency_results.get("issues", []):
            if ci["type"] == "value_inconsistency":
                issues.append({
                    "category": "consistency",
                    "action": "fix_value_inconsistency",
                    "detail": f"書類間の数値不整合: {ci['description']}",
                })
            elif ci["type"] == "negative_profit":
                issues.append({
                    "category": "consistency",
                    "action": "fix_negative_profit",
                    "detail": f"営業利益がマイナス: {ci['description']}",
                })
    else:
        consistency_score = 0
    breakdown["consistency"] = {"score": round(consistency_score, 1), "max": 25, "detail": f"整合率: {consistency_results.get('score_ratio', 0):.0%}"}
    score += consistency_score

    return {
        "score": round(score, 1),
        "max": 100,
        "breakdown": breakdown,
        "issues": issues,
        "raw": {
            "files": file_results,
            "diagrams": diagram_results,
            "text": text_results,
            "values": value_results,
            "text_quality": quality_results,
            "consistency": consistency_results,
        },
    }


# =============================================================================
# レポート出力
# =============================================================================

def print_report(file_results, diagram_results, text_results, value_results):
    """検証レポートを出力"""
    print("=" * 60)
    print("省力化補助金 申請書類 検証レポート")
    print("=" * 60)

    # ファイル存在チェック
    print("\n--- ファイル存在チェック ---")
    file_ok = 0
    for r in file_results:
        status = "OK" if r["ok"] else "NG"
        size_str = f"({r['size']:,}B)" if r["exists"] else ""
        print(f"  [{status}] {r['file']} {size_str}")
        if r["ok"]:
            file_ok += 1
    print(f"  結果: {file_ok}/{len(file_results)} ファイル")

    # 図解チェック
    print("\n--- 図解チェック ---")
    if "note" in diagram_results:
        print(f"  [SKIP] {diagram_results['note']}")
    else:
        status = "OK" if diagram_results["ok"] else "NG"
        print(f"  [{status}] 図解: {diagram_results['found']}/{diagram_results['expected']}枚")

    # 文字数チェック
    print("\n--- 文字数チェック ---")
    if "error" in text_results:
        print(f"  [SKIP] {text_results['error']}")
    else:
        status = "OK" if text_results["ok"] else "NG"
        print(f"  [{status}] 総文字数: {text_results['total_chars']:,}字 (基準: {text_results['min_required']:,}字以上)")

        # セクション別
        sections = text_results.get("sections", {})
        if sections:
            print("  --- セクション別 ---")
            for key in sorted(sections.keys()):
                sec = sections[key]
                s = "OK" if sec["ok"] else "NG"
                print(f"    [{s}] {key}: {sec['chars']:,}字 (基準: {sec['min_required']:,}字以上)")

    # 数値チェック
    print("\n--- 数値整合性チェック ---")
    if isinstance(value_results, dict) and "error" in value_results:
        print(f"  [SKIP] {value_results['error']}")
    else:
        for key, val in value_results.items():
            if isinstance(val, dict) and "ok" in val:
                status = "OK" if val["ok"] else "NG"
                detail = " / ".join(f"{k}: {v}" for k, v in val.items() if k != "ok")
                print(f"  [{status}] {key}: {detail}")

    # 総合判定
    print("\n" + "=" * 60)
    all_files_ok = all(r["ok"] for r in file_results)
    diagrams_ok = diagram_results.get("ok", False)
    text_ok = text_results.get("ok", False) if "error" not in text_results else True
    sections_ok = all(
        s.get("ok", True) for s in text_results.get("sections", {}).values()
    ) if "error" not in text_results else True
    values_ok = all(
        v.get("ok", True) for v in value_results.values()
        if isinstance(v, dict) and "ok" in v
    ) if isinstance(value_results, dict) and "error" not in value_results else True

    if all_files_ok and diagrams_ok and text_ok and sections_ok and values_ok:
        print("総合判定: PASS")
        return 0
    else:
        issues = []
        if not all_files_ok:
            issues.append(f"ファイル不足({file_ok}/{len(file_results)})")
        if not diagrams_ok:
            issues.append(f"図解不足({diagram_results.get('found', 0)}/{EXPECTED_DIAGRAM_COUNT})")
        if not text_ok:
            issues.append(f"総文字数不足({text_results.get('total_chars', 0)}字)")
        if not sections_ok:
            ng_sections = [k for k, v in text_results.get("sections", {}).items() if not v.get("ok")]
            issues.append(f"セクション文字数不足({', '.join(ng_sections)})")
        if not values_ok:
            issues.append("数値要件未達")
        print(f"総合判定: FAIL ({', '.join(issues)})")
        return 1


def build_json_report(file_results, diagram_results, text_results, value_results):
    """JSON形式のレポートを構築"""
    return {
        "files": file_results,
        "diagrams": diagram_results,
        "text": text_results,
        "values": value_results,
        "pass": (
            all(r["ok"] for r in file_results)
            and diagram_results.get("ok", False)
            and text_results.get("ok", True)
            and all(
                v.get("ok", True) for v in value_results.values()
                if isinstance(v, dict) and "ok" in v
            ) if isinstance(value_results, dict) and "error" not in value_results else True
        ),
    }


# =============================================================================
# メイン
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="省力化補助金 申請書類 出力検証")
    parser.add_argument("--output", "-o", required=True, help="出力ディレクトリ")
    parser.add_argument("--json", action="store_true", help="JSON形式で出力")
    args = parser.parse_args()

    output_dir = Path(args.output)
    if not output_dir.exists():
        print(f"エラー: ディレクトリが存在しない: {output_dir}")
        sys.exit(1)

    file_results = check_files(output_dir)
    diagram_results = check_diagrams(output_dir)
    text_results = check_docx_text(output_dir)
    value_results = check_plan3_values(output_dir)

    if args.json:
        report = build_json_report(file_results, diagram_results, text_results, value_results)
        print(json.dumps(report, ensure_ascii=False, indent=2))
        sys.exit(0 if report["pass"] else 1)
    else:
        exit_code = print_report(file_results, diagram_results, text_results, value_results)
        sys.exit(exit_code)


if __name__ == "__main__":
    main()
