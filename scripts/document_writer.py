#!/usr/bin/env python3
"""事業計画書Part1-2 Word文書生成"""

import os
import shutil
from pathlib import Path
from typing import Dict

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import HearingData
from config import Config
from content_generator import ContentGenerator
from hearing_reader import _find_sheet_in_workbook


def generate_business_plan_1_2(data: HearingData, diagrams: Dict[str, str], output_dir: str, template_path: Path):
    """事業計画書その1その2を生成"""
    print("\n📝 事業計画書（その1＋その2）を生成中...")

    output_path = Path(output_dir) / "事業計画書_その1その2_完成版.docx"
    shutil.copy(template_path, output_path)
    os.chmod(output_path, 0o644)

    doc = Document(output_path)
    gen = ContentGenerator(data)
    c, s, l, e, f = data.company, data.labor_shortage, data.labor_saving, data.equipment, data.funding

    manufacturer = e.manufacturer if e.manufacturer else "オーダーメイド開発"
    model_name = e.model if e.model else "カスタム仕様"

    # ヘルパー関数
    def get_unique_cells(row):
        unique, seen = [], set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen:
                seen.add(cid)
                unique.append(cell)
        return unique

    def clear_and_write(cell, text):
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""
        if cell.paragraphs:
            cell.paragraphs[0].text = text
        else:
            cell.text = text

    # ----- テーブル0: 事業者情報 -----
    print("    📋 事業者情報...")
    if len(doc.tables) > 0:
        t = doc.tables[0]
        info = [c.name, f"代表取締役  {c.representative}", f"{c.prefecture}{c.address}",
                c.industry, c.established_date, f"{c.officer_count}名 ／ {c.employee_count}名", c.url or ""]
        for i, val in enumerate(info):
            if i < len(t.rows) and len(t.rows[i].cells) > 1:
                t.rows[i].cells[1].text = val

    # ----- テーブル1: 事業計画名 -----
    print("    📋 事業計画名...")
    if len(doc.tables) > 1:
        doc.tables[1].rows[0].cells[0].text = f"{e.name}の導入による業務省力化と生産性向上"[:30]

    # ----- テーブル2: 概要 -----
    print("    📋 事業計画概要...")
    if len(doc.tables) > 2:
        doc.tables[2].rows[0].cells[0].text = f"当社は{c.industry}を営む企業である。{s.shortage_tasks}において人手不足が深刻であり、月{s.overtime_hours}時間の残業が発生している。{e.name}を導入し、作業時間を{l.reduction_rate:.0f}%削減することで、生産性向上と従業員の負担軽減を実現する。"

    # ----- テーブル3: 導入設備 -----
    print("    📋 導入設備情報...")
    if len(doc.tables) > 3:
        doc.tables[3].rows[0].cells[0].text = f"【設備名称】{e.name}\n【メーカー】{manufacturer}\n【型番】{model_name}\n【数量】{e.quantity}台\n【金額】{e.total_price:,}円（税抜）\n【購入先】{e.vendor}"

    # ----- テーブル4: ネストテーブル + 本文 -----
    if len(doc.tables) > 4:
        t4 = doc.tables[4]

        # ネストテーブル（事業者概要）
        print("    📋 事業者概要テーブル（ネスト）...")
        cell0 = t4.rows[0].cells[0]
        if cell0.tables:
            nested = cell0.tables[0]
            overview = gen.generate_business_overview_table_data()

            # 行0-5: テキスト項目
            text_items = ["事業者名", "経営理念", "経営戦略", "事業コンセプト", "事業内容", "長期的なビジョン"]
            for row_idx, key in enumerate(text_items):
                if row_idx < len(nested.rows):
                    uc = get_unique_cells(nested.rows[row_idx])
                    if len(uc) >= 2:
                        clear_and_write(uc[1], overview[key])

            # 行7-10: 直近実績
            fin_data = overview["直近実績"]
            fin_rows = [(7, "売上金額"), (8, "売上総利益"), (9, "営業利益"), (10, "従業員数")]
            for row_idx, key in fin_rows:
                if row_idx < len(nested.rows):
                    uc = get_unique_cells(nested.rows[row_idx])
                    if len(uc) >= 4:
                        vals = fin_data[key]
                        fmt = lambda v: f"{v:,}円" if key != "従業員数" else f"{v}名"
                        clear_and_write(uc[1], fmt(vals[0]))
                        clear_and_write(uc[2], fmt(vals[1]))
                        clear_and_write(uc[3], fmt(vals[2]))

        # 本文セクション
        print("    📋 本文セクション（PREP法）...")
        sections = {
            1: gen.generate_section_1_1() + "\n\n" + gen.generate_swot_analysis(),
            2: gen.generate_section_1_2(),
            3: gen.generate_section_1_3(),
            4: f"【導入設備の詳細】\n設備名称：{e.name}\nメーカー：{manufacturer}\n型番：{model_name}\n数量：{e.quantity}台\n金額：{e.total_price:,}円（税抜）\n購入先：{e.vendor}\nカタログ番号：{e.catalog_number or 'オーダーメイド'}\n\n【設備の特徴】\n{e.features}\n\n【投資金額の内訳】\n事業費総額：{f.total_investment:,}円\n補助金申請額：{f.subsidy_amount:,}円\n自己負担額：{f.self_funding:,}円",
            5: gen.generate_section_2_1(),
            6: gen.generate_section_2_2(),
            8: gen.generate_section_3_1(),
            9: f"【資金調達計画】\n事業費総額：{f.total_investment:,}円\nうち補助金：{f.subsidy_amount:,}円\nうち自己資金：{f.self_funding:,}円\n\n自己資金については、当社の内部留保および取引銀行である{f.bank_name}からの借入により調達する予定である。\n\n【投資回収計画】\n本設備への投資は、省力化による人件費削減効果と売上拡大による利益増加により、約2〜3年での回収を見込んでいる。",
            10: f"【実施体制】\n統括責任者：{c.representative}（代表取締役）\n実施責任者：{f.implementation_manager}\n従業員{c.employee_count}名と連携して実施\n\n【スケジュール】\n実施期間：{f.implementation_period}\n\n1ヶ月目：契約・発注\n2ヶ月目：設備納品・設置工事\n3ヶ月目：試運転・調整・従業員教育\n4ヶ月目以降：本格稼働・効果測定",
            11: f"【人手不足の状況】\n当社は「限られた人手で業務を遂行するため、直近の従業員の平均残業時間が30時間を超えている」状況に該当する。直近12ヶ月の平均残業時間：月{s.overtime_hours}時間\n\n【オーダーメイド性】\n本設備は当社の業務に特化したカスタマイズを施す。{e.features}\n\n【賃上げ計画の表明】\n・1人当たり給与支給総額の年平均成長率：{(Config.SALARY_GROWTH_RATE - 1) * 100:.1f}%以上\n・事業場内最低賃金：{c.prefecture}の地域別最低賃金を30円以上上回る水準"
        }

        for row_idx, content in sections.items():
            if row_idx < len(t4.rows):
                cell = t4.rows[row_idx].cells[0]
                existing = cell.text
                cell.text = existing.rstrip() + "\n\n" + content.strip()
                print(f"      ✅ セクション{row_idx}（{len(content)}文字）")

        # 図解挿入
        if diagrams:
            print("    🖼️ 図解挿入...")
            mapping = {1: ["01_企業概要", "02_SWOT分析"], 2: ["03_人手不足", "04_課題フロー", "12_業務フロー"],
                       4: ["05_設備概要"], 5: ["06_ビフォーアフター", "13_工程別比較"], 6: ["07_効果算定"],
                       10: ["08_実施体制", "09_スケジュール"], 8: ["10_5年計画"]}
            for row_idx, ids in mapping.items():
                if row_idx < len(t4.rows):
                    cell = t4.rows[row_idx].cells[0]
                    for did in ids:
                        if did in diagrams:
                            try:
                                para = cell.add_paragraph()
                                para.add_run().add_picture(diagrams[did], width=Inches(5.5))
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as ex:
                                print(f"      ⚠️ 図解挿入エラー ({did}): {ex}")

    # スケジュール図を追加（11_実施工程）
    if "11_実施工程" in diagrams:
        print("    📅 補助事業スケジュール図...")
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run("補助事業のスケジュール（参考）")
        run.bold = True
        run.font.size = Pt(14)
        p = doc.add_paragraph()
        try:
            p.add_run().add_picture(diagrams["11_実施工程"], width=Inches(6.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as ex:
            print(f"      ⚠️ スケジュール図挿入エラー: {ex}")

    doc.save(output_path)
    print(f"  ✅ 保存完了: {output_path}")


def add_schedule_table(doc, data: HearingData):
    """補助事業のスケジュール表をWord表形式で追加"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import RGBColor

    base_year = 2026  # 交付決定想定年度

    # ページ区切り
    doc.add_page_break()

    # タイトル
    p = doc.add_paragraph()
    run = p.add_run("補助事業のスケジュール（参考）")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run(
        "前述の補助事業の内容に沿い機械装置等の取得時期や技術の導入時期を含めた"
        "スケジュールを示してください。記載例ですので適宜使いやすいように作成してください。"
    )
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run("≪事業計画期間にかかるスケジュール≫")
    run.bold = True
    run.font.size = Pt(11)

    # フェーズ定義
    phases = [
        ("0．構想設計", [
            ("事業目的・目標設定", list(range(0, 3)), []),
            ("課題・改善方針検討", list(range(0, 4)), []),
            ("事業計画作成", list(range(1, 5)), []),
            ("社内プロジェクト体\n制決定", list(range(1, 4)), []),
            ("投資採算性・投資規\n模決定", list(range(2, 6)), []),
            ("予算・調達計画策定", list(range(3, 6)), []),
        ]),
        ("1．機能設計", [
            ("システム要件定義", list(range(3, 6)), []),
            ("システム構成策定", list(range(4, 7)), []),
            ("機能一覧定義", list(range(5, 8)), []),
        ]),
        ("2．周辺機器の手配", [
            ("機械装置発注", list(range(5, 7)), []),
            ("部品・原材料調達", list(range(5, 9)), []),
        ]),
        ("3．機能試作、シス\nテム組み立て", [
            ("システム設計", list(range(6, 9)), []),
            ("システム発注・開発", list(range(7, 10)), []),
        ]),
        ("4．評価", [
            ("テスト・リリース", list(range(8, 10)), []),
            ("課題・改善方針検討", list(range(9, 11)), []),
        ]),
        ("5．調整改善", [
            ("システム再設計", list(range(10, 12)), []),
        ]),
        ("6．稼働・実装", [
            ("セキュリティ対策", list(range(11, 13)), [0, 1]),
            ("保守・管理", [12], [0, 1, 2, 3, 4]),
        ]),
    ]

    total_tasks = sum(len(tasks) for _, tasks in phases)
    HEADER_ROWS = 3
    TOTAL_COLS = 20  # 2(フェーズ+タスク) + 13(月) + 5(年)

    table = doc.add_table(rows=HEADER_ROWS + total_tasks, cols=TOTAL_COLS)
    table.style = 'Table Grid'

    def shade_cell(cell, color="B4C6E7"):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell(cell, text, size=7, bold=False, align=None, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        if align:
            p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    def mark_active(cell, color="B4C6E7"):
        """活動期間セルにマーカーと背景色を設定"""
        shade_cell(cell, color)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("⇨")
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0x30, 0x60, 0xA0)

    # --- ヘッダー行0: 大見出し ---
    table.cell(0, 0).merge(table.cell(0, 1))
    set_cell(table.cell(0, 0), "", 7)
    table.cell(0, 2).merge(table.cell(0, 14))
    set_cell(table.cell(0, 2), "補助事業実施期間", 8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    table.cell(0, 15).merge(table.cell(0, 19))
    set_cell(table.cell(0, 15), "事業計画期間", 8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- ヘッダー行1: 年度ラベル ---
    table.cell(1, 0).merge(table.cell(1, 1))
    set_cell(table.cell(1, 0), "", 7)
    table.cell(1, 2).merge(table.cell(1, 14))
    set_cell(table.cell(1, 2), "", 7)
    for i in range(5):
        set_cell(table.cell(1, 15 + i), f"事業計画{i+1}年目", 6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- ヘッダー行2: 月＋年度期間 ---
    set_cell(table.cell(2, 0), "", 7)
    set_cell(table.cell(2, 1), "", 7)
    month_labels = ["3\n月", "4\n月", "5\n月", "6\n月", "7\n月", "8\n月",
                    "9\n月", "10\n月", "11\n月", "12\n月", "1", "2", "3"]
    for i, label in enumerate(month_labels):
        set_cell(table.cell(2, 2 + i), label, 6, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(5):
        ys = base_year + i
        ye = ys + 1
        if i == 0:
            label = f"※{ys}年4月～\n{ye}年3月"
        else:
            label = f"{ys}年4月～{ye}\n年3月"
        set_cell(table.cell(2, 15 + i), label, 5, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- データ行 ---
    current_row = HEADER_ROWS
    for phase_name, tasks in phases:
        start_row = current_row
        for task_name, active_months, active_years in tasks:
            row = current_row
            set_cell(table.cell(row, 1), task_name, 7)
            for m in active_months:
                if 0 <= m <= 12:
                    mark_active(table.cell(row, 2 + m))
            for y in active_years:
                if 0 <= y <= 4:
                    mark_active(table.cell(row, 15 + y))
            current_row += 1

        end_row = current_row - 1
        if start_row < end_row:
            table.cell(start_row, 0).merge(table.cell(end_row, 0))
        set_cell(table.cell(start_row, 0), phase_name, 7)

    # フッター注記
    p = doc.add_paragraph()
    run = p.add_run("事業計画は事業者ごとの決算期")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
